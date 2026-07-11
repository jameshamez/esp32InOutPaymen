#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_installation_docx import (
    BLACK,
    BLUE,
    DARK_BLUE,
    GREEN,
    MUTED,
    NAVY,
    add_body,
    add_callout,
    add_code_block,
    add_heading,
    add_page_field,
    add_table,
    configure_document,
    new_abstract_numbering,
    new_number_instance,
    set_cell_fill,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
PACKAGE_NAME = "PaymentESP-v1.1.0-Hardware-20260628"
OUTPUT = ROOT / "deliverables" / PACKAGE_NAME / "documents" / "PaymentESP-Hardware-Test-Report-TH.docx"
VERSION = "1.1.0-Hardware"
REPORT_DATE = "28 มิถุนายน 2026"
FIRMWARE_SHA256 = "14b1c98f10fffd7c5dd8dbed3a100f7b373b63836dd449cb72ba0e80b8bbfc2f"
MERGED_SHA256 = "637e6919738f92bf1d254b2e34826aee8a3c83024ac523c96af025ee9957cd9a"


def configure_report(doc):
    """Resolve the standard_business_brief preset for this test report."""
    configure_document(doc)
    section = doc.sections[0]
    section.different_first_page_header_footer = False

    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.clear()
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("PaymentESP | รายงานผลการทดสอบ Hardware Edition"), size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run("หน้า "), size=9, color=MUTED)
    add_page_field(footer, "PAGE")


def add_report_table(doc, headers, rows, widths, font_size=9.3):
    table = add_table(doc, headers, rows, widths, font_size=font_size)
    for cell in table.rows[0].cells:
        set_cell_fill(cell, "F2F4F7")
    return table


def add_report_list(doc, items, ordered=False):
    abstract_id = new_abstract_numbering(
        doc,
        "decimal" if ordered else "bullet",
        "%1." if ordered else "•",
        left=720,
        hanging=360,
    )
    num_id = new_number_instance(doc, abstract_id)
    for item in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.167
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_el])
        paragraph._p.get_or_add_pPr().insert(0, num_pr)
        set_run_font(paragraph.add_run(item))


def add_masthead(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(20)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    set_run_font(kicker.add_run("TEST REPORT"), size=10.5, color=GREEN, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    set_run_font(title.add_run("รายงานผลการทดสอบ PaymentESP"), size=24, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("ESP32 PromptPay QR Payment | Hardware Delivery Edition"), size=13, color=MUTED)

    metadata = [
        ("โครงการ", "ระบบ QR Payment PromptPay ด้วย ESP32"),
        ("เวอร์ชัน", VERSION),
        ("วันที่ทดสอบ", REPORT_DATE),
        ("สภาพแวดล้อม", "Host unit/integration tests, PlatformIO esp32dev และ Wokwi build"),
        ("สถานะ", "Automated Test PASS | Hardware Acceptance PENDING"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        set_run_font(paragraph.add_run(f"{label}: "), size=10.5, color=BLACK, bold=True)
        set_run_font(paragraph.add_run(value), size=10.5, color=BLACK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(14)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GREEN)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    add_callout(
        doc,
        "สรุปผล",
        "ชุดทดสอบอัตโนมัติและการ Build ผ่านทั้งหมด ไม่พบ Fail ในรอบนี้ ส่วนการ Upload, OLED, WiFi หน้างาน และการสแกนด้วยแอปธนาคารยังต้องตรวจบน ESP32 เครื่องจริงก่อนส่งลูกค้า",
        "success",
    )


def build_document():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_report(doc)

    props = doc.core_properties
    props.title = "รายงานผลการทดสอบ PaymentESP Hardware Edition"
    props.subject = "ESP32 PromptPay QR automated test and hardware acceptance report"
    props.author = "PaymentESP Project Team"
    props.keywords = "ESP32, PromptPay, Test Report, PlatformIO, Dashboard, REST POST"
    props.comments = f"Test report version {VERSION}"

    add_masthead(doc)

    add_heading(doc, "1. วัตถุประสงค์และขอบเขต", 1)
    add_body(doc, "รายงานฉบับนี้บันทึกผลที่ทดสอบซ้ำจาก Source Code และ Firmware ชุดส่งมอบ โดยแยกผลทดสอบอัตโนมัติออกจาก Hardware Acceptance อย่างชัดเจน เพื่อใช้เป็นหลักฐานก่อนนำ Firmware ไปติดตั้งบน ESP32 ตัวจริง")
    add_report_list(doc, [
        "ตรวจความถูกต้องของ PromptPay payload, CRC, Static QR และ Dynamic QR",
        "ตรวจ Local Dashboard, การบันทึกค่า, validation, webhook และ logs",
        "ตรวจ Omise module ในระดับ mock integration เท่านั้น ไม่ได้เรียก Live/Test API ภายนอก",
        "Build Firmware สำหรับ ESP32 DevKit V1 และ Wokwi environment",
        "ระบุรายการที่ต้องตรวจเพิ่มเมื่อมี ESP32/OLED เครื่องจริงต่อกับเครื่องพัฒนา",
    ])
    add_callout(doc, "ขอบเขตสำคัญ", "HTTP 202 จาก REST receiver หมายถึงรับเหตุการณ์สร้าง QR สำเร็จ ไม่ใช่หลักฐานว่าเงินเข้าบัญชีแล้ว", "warning")

    add_heading(doc, "2. สภาพแวดล้อมทดสอบ", 1)
    add_report_table(doc, ["รายการ", "รายละเอียด"], [
        ("วันที่", REPORT_DATE),
        ("Firmware framework", "Arduino บน PlatformIO"),
        ("PlatformIO Core", "6.1.19"),
        ("Hardware target", "ESP32 DevKit V1, environment esp32dev"),
        ("Simulation target", "Wokwi environment สำหรับ compile/simulation เสริม"),
        ("Dashboard", "Node.js local integration test บน localhost และ temporary storage"),
        ("Secret/PromptPay จริง", "ไม่ใช้และไม่บันทึกในรายงานฉบับนี้"),
        ("ESP32 serial device", "ไม่พบอุปกรณ์ USB Serial ระหว่างจัดทำรายงาน"),
    ], [2800, 6560], font_size=9.5)

    add_heading(doc, "3. สรุปผลการทดสอบ", 1)
    add_report_table(doc, ["กลุ่มทดสอบ", "วิธี", "ผล"], [
        ("PromptPay QR", "C++ host unit test", "PASS"),
        ("Local Dashboard", "Node.js integration test", "PASS"),
        ("Omise module", "Mock integration test", "PASS"),
        ("ESP32 Hardware Build", "pio run -e esp32dev", "PASS"),
        ("Wokwi Build", "pio run -e wokwi", "PASS"),
        ("Physical Hardware Acceptance", "Upload และตรวจ ESP32/OLED จริง", "PENDING"),
    ], [3900, 3660, 1800], font_size=9.3)
    add_callout(doc, "ผลรวม", "Automated/Build 5 กลุ่มผ่านทั้งหมด, Fail 0 กลุ่ม และ Physical Hardware Acceptance รอตรวจ", "note")

    doc.add_page_break()
    add_heading(doc, "4. Test Case Matrix", 1)
    test_rows = [
        ("TR-01", "CRC-16/CCITT-FALSE known vector", "ค่า 123456789 ต้องได้ 0x29B1", "PASS"),
        ("TR-02", "Static QR สำหรับเบอร์โทร", "Tag Point of Initiation เป็น 11, ไม่มี Amount, CRC ถูกต้อง", "PASS"),
        ("TR-03", "Dynamic QR พร้อมยอดเงิน", "Tag เป็น 12, Amount/Reference อยู่ใน payload, CRC ถูกต้อง", "PASS"),
        ("TR-04", "PromptPay National ID", "จำแนก ID 13 หลักและสร้าง payload ถูกต้อง", "PASS"),
        ("TR-05", "Event JSON", "เวลา, epoch และ webhookStatus serialize ได้", "PASS"),
        ("TR-06", "Dashboard config persistence", "บันทึก PromptPay/Webhook ลงไฟล์ชั่วคราวและอ่านคืนได้", "PASS"),
        ("TR-07", "Dynamic amount validation", "15 แปลงเป็น 15.00 และยอด 0 ถูกปฏิเสธ", "PASS"),
        ("TR-08", "Static amount behavior", "Static QR ไม่ใส่ Amount แม้ request ส่งยอดมา", "PASS"),
        ("TR-09", "Webhook receiver", "POST event ได้ HTTP 202 และบันทึก reference", "PASS"),
        ("TR-10", "Invalid configuration", "PromptPay/Webhook รูปแบบผิดถูกตอบ 400", "PASS"),
        ("TR-11", "Omise normalize", "แปลง amount/reference/QR URL จาก mock response", "PASS"),
        ("TR-12", "Omise request", "สร้าง PromptPay charge request และ retrieve charge ผ่าน mock fetch", "PASS"),
        ("TR-13", "Omise input security", "ปฏิเสธ charge ID ที่มี path traversal", "PASS"),
        ("TR-14", "ESP32 Hardware build", "Compile environment esp32dev สำเร็จ", "PASS"),
        ("TR-15", "Wokwi build", "Compile environment wokwi สำเร็จ", "PASS"),
    ]
    add_report_table(doc, ["ID", "รายการ", "เกณฑ์ผ่าน", "ผล"], test_rows, [1100, 2870, 3890, 1500], font_size=8.2)

    add_heading(doc, "5. คำสั่งและผลที่ยืนยัน", 1)
    add_heading(doc, "5.1 PromptPay unit test", 2)
    add_code_block(doc, "c++ -std=c++17 -Iinclude src/PromptPayQR.cpp tests/test_promptpay.cpp -o test_promptpay\n./test_promptpay\n# PromptPay QR tests passed")
    add_heading(doc, "5.2 Dashboard และ Omise mock integration", 2)
    add_code_block(doc, "node tests/test_dashboard.js\n# Dashboard integration tests passed\n\nnode tests/test_omise.js\n# Omise integration tests passed")
    add_callout(doc, "หมายเหตุ Omise", "การทดสอบนี้ใช้ mock response และ mock fetch จึงตรวจ logic ภายในเท่านั้น ไม่ยืนยันการเชื่อมต่อบัญชี Omise จริง", "note")

    add_heading(doc, "5.3 PlatformIO build", 2)
    add_code_block(doc, "pio run -e esp32dev\n# SUCCESS\n\npio run -e wokwi\n# SUCCESS")
    add_report_table(doc, ["Environment", "RAM", "Flash", "ผล"], [
        ("esp32dev", "50,408 / 327,680 (15.4%)", "1,052,917 / 1,310,720 (80.3%)", "PASS"),
        ("wokwi", "50,408 / 327,680 (15.4%)", "1,052,973 / 1,310,720 (80.3%)", "PASS"),
    ], [1900, 2750, 3210, 1500], font_size=8.7)

    add_heading(doc, "6. Firmware Integrity", 1)
    add_report_table(doc, ["ไฟล์", "SHA-256"], [
        ("firmware.bin", FIRMWARE_SHA256),
        ("firmware-merged.bin", MERGED_SHA256),
    ], [2500, 6860], font_size=8.3)
    add_body(doc, "ค่า Hash ใช้ตรวจว่าไฟล์ Firmware ที่นำไปติดตั้งตรงกับไฟล์ที่ Build และจัดส่งในรอบรายงานนี้")

    doc.add_page_break()
    add_heading(doc, "7. Requirement Traceability", 1)
    add_report_table(doc, ["Requirement", "หลักฐานทดสอบ", "สถานะ"], [
        ("WiFi และตั้งค่า PromptPay ID", "Source path/compile ผ่าน; Setup AP และ NVS ต้องตรวจบนบอร์ดจริง", "PARTIAL"),
        ("Static QR PromptPay", "TR-02 และ TR-08", "PASS"),
        ("Dynamic QR ระบุยอดเงิน", "TR-03 และ TR-07", "PASS"),
        ("Serial Output", "Source/compile ผ่าน; output จาก USB Serial จริงรอตรวจ", "PARTIAL"),
        ("REST API POST", "TR-09", "PASS"),
        ("Logging เวลาและยอดเงิน", "TR-05, TR-06 และ TR-09", "PASS"),
        ("แสดง QR บน OLED", "Build ผ่าน; การแสดงผลบน OLED จริงรอตรวจ", "PENDING"),
        ("ตรวจเงินเข้าบัญชีอัตโนมัติ", "อยู่นอกขอบเขตใบเสนอราคาและต้องใช้ Bank/Gateway API", "OUT OF SCOPE"),
    ], [3260, 4460, 1640], font_size=8.7)

    doc.add_page_break()
    add_heading(doc, "8. Hardware Acceptance ที่ยังต้องทำ", 1)
    add_callout(doc, "สถานะปัจจุบัน", "ไม่พบบอร์ด ESP32 USB Serial ต่ออยู่กับเครื่องขณะจัดทำรายงาน จึงไม่ระบุผล Hardware เป็น PASS", "warning")
    acceptance_rows = [
        ("HA-01", "Upload firmware ลง ESP32 ตัวที่จะส่ง", "PENDING"),
        ("HA-02", "OLED SSD1306 แสดงข้อความ Setup/Ready และไม่ค้าง", "PENDING"),
        ("HA-03", "เปิด AP PaymentESP-Setup และเข้า 192.168.4.1/setup", "PENDING"),
        ("HA-04", "บันทึก WiFi/PromptPay/Webhook และรีสตาร์ตได้", "PENDING"),
        ("HA-05", "ปิดเปิดไฟแล้วค่าจาก NVS ยังคงอยู่", "PENDING"),
        ("HA-06", "เปิด paymentesp.local หรือ IP จากอุปกรณ์ใน LAN", "PENDING"),
        ("HA-07", "Static QR สแกนด้วยแอปธนาคารและชื่อผู้รับถูกต้อง", "PENDING"),
        ("HA-08", "Dynamic QR 15.00 บาทสแกนและแสดงยอดถูกต้อง", "PENDING"),
        ("HA-09", "Serial, Logs และ REST POST แสดง amount/reference ตรงกัน", "PENDING"),
        ("HA-10", "ถ่ายภาพ OLED, Serial, Logs และ REST เป็นหลักฐาน", "PENDING"),
    ]
    add_report_table(doc, ["ID", "รายการตรวจ", "สถานะ"], acceptance_rows, [1300, 6260, 1800], font_size=9)

    add_heading(doc, "9. เกณฑ์ปิดงาน", 1)
    add_report_list(doc, [
        "เปลี่ยน HA-01 ถึง HA-10 เป็น PASS หลังทดสอบกับเครื่องจริง",
        "แนบภาพหลักฐานที่ไม่เปิดเผย WiFi Password, Secret Key หรือ PromptPay payload เต็ม",
        "บันทึก Serial Port, หมายเลขบอร์ด/Asset ID และวันที่ตรวจรับ",
        "ให้ผู้ทดสอบและผู้ส่งมอบลงชื่อก่อนบรรจุเครื่อง",
    ], ordered=True)

    doc.add_page_break()
    add_heading(doc, "10. ลงนามผลการตรวจรับ Hardware", 1)
    add_report_table(doc, ["บทบาท", "ชื่อ/ลายเซ็น", "วันที่", "ผล"], [
        ("ผู้ทดสอบ", "", "", "PASS / FAIL"),
        ("ผู้ส่งมอบ", "", "", "ยืนยันส่งมอบ"),
        ("ผู้รับมอบ", "", "", "ยอมรับ / มีข้อแก้ไข"),
    ], [2000, 3160, 1900, 2300], font_size=9.5)
    add_callout(doc, "ข้อสรุป", "เอกสารฉบับนี้รับรองเฉพาะผล Automated Test และ Build ที่ระบุไว้ การรับรอง Hardware จะสมบูรณ์เมื่อ HA-01 ถึง HA-10 ได้รับการตรวจและลงนาม", "note")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
