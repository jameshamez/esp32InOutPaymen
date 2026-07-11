#!/usr/bin/env python3
from pathlib import Path
import sys

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "PaymentESP-v1.0.0-20260628" / "documents" / "PaymentESP-Detailed-Installation-Manual-TH.docx"
ASSET_DIR = ROOT / "docs" / "word-assets"
EVIDENCE_DIR = ROOT / "docs" / "evidence"
TABLE_HELPER = Path("/Users/nutthapongphonruk/.codex/plugins/cache/openai-primary-runtime/documents/26.623.12021/skills/documents/scripts")
sys.path.insert(0, str(TABLE_HELPER))
from table_geometry import apply_table_geometry  # noqa: E402


FONT = "Sarabun"
MONO_FONT = "Menlo"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17324D"
MUTED = "5B6878"
GREEN = "08766B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "EAF8F4"
LIGHT_GRAY = "F4F6F9"
CAUTION = "FFF4D6"
RED_LIGHT = "FDECEC"
WHITE = "FFFFFF"
BLACK = "182230"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 100, "bottom": 100, "start": 120, "end": 120}


def set_run_font(run, name=FONT, size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:cantSplit")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_border(cell, color="D7DEE8", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, color):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_paragraph_border(paragraph, color="C8D2DE", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), color)
        borders.append(border)


def add_page_field(paragraph, field_name):
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)
    set_run_font(begin_run, size=9, color=MUTED)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    instr_run._r.append(instr)
    set_run_font(instr_run, size=9, color=MUTED)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    set_run_font(separate_run, size=9, color=MUTED)

    result_run = paragraph.add_run("1")
    set_run_font(result_run, size=9, color=MUTED)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    set_run_font(end_run, size=9, color=MUTED)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    return paragraph, run


def add_body(doc, text, bold_prefix=None, italic=False):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, italic=italic)
    return paragraph


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    set_paragraph_shading(paragraph, "EEF2F6")
    run = paragraph.add_run(lines)
    set_run_font(run, name=MONO_FONT, size=8.6, color="263244")
    return paragraph


def add_callout(doc, label, text, kind="note"):
    fills = {"note": LIGHT_BLUE, "success": LIGHT_GREEN, "warning": CAUTION, "danger": RED_LIGHT}
    colors = {"note": DARK_BLUE, "success": GREEN, "warning": "7A5A00", "danger": "9B1C1C"}
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(paragraph, fills[kind])
    set_paragraph_border(paragraph, color="C8D2DE", size="6")
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, color=colors[kind], bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=BLACK)
    return paragraph


def new_abstract_numbering(doc, fmt, text, left=540, hanging=270):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    level.append(lvl_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(left))
    indent.set(qn("w:hanging"), str(hanging))
    p_pr.extend([tabs, indent])
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def new_number_instance(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list(doc, items, ordered=False):
    abstract_id = new_abstract_numbering(doc, "decimal" if ordered else "bullet", "%1." if ordered else "•")
    num_id = new_number_instance(doc, abstract_id)
    paragraphs = []
    for item in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_el])
        p_pr.insert(0, num_pr)
        run = paragraph.add_run(item)
        set_run_font(run)
        paragraphs.append(paragraph)
    return paragraphs


def add_table(doc, headers, rows, widths, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_fill(cell, LIGHT_BLUE)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=font_size, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cell = cells[index]
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            if index == 0 and len(headers) <= 3:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size)
        set_row_cant_split(table.rows[-1])
    apply_table_geometry(table, widths, table_width_dxa=CONTENT_DXA, indent_dxa=TABLE_INDENT_DXA, cell_margins_dxa=CELL_MARGINS)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, path, caption, width=6.1):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(2)
    caption_p.paragraph_format.space_after = Pt(10)
    caption_run = caption_p.add_run(caption)
    set_run_font(caption_run, size=9, color=MUTED, italic=True)


def build_architecture_diagram(output):
    width, height = 1600, 650
    image = Image.new("RGB", (width, height), "#F5F7FA")
    draw = ImageDraw.Draw(image)
    font_path = "/System/Library/AssetsV2/com_apple_MobileAsset_Font8/cf0dc8d3b09f9ba379660e591e82566e2b557949.asset/AssetData/Sarabun.ttc"
    title_font = ImageFont.truetype(font_path, 42)
    body_font = ImageFont.truetype(font_path, 31)
    small_font = ImageFont.truetype(font_path, 24)
    draw.text((60, 35), "โครงสร้างการทำงาน PaymentESP", fill="#17324D", font=title_font)

    boxes = {
        "dashboard": (70, 180, 390, 410, "Local Dashboard\nlocalhost:3001", "#E8EEF5", "#2E74B5"),
        "esp32": (640, 155, 985, 435, "ESP32\nPromptPay QR API", "#EAF8F4", "#08766B"),
        "oled": (1210, 85, 1510, 285, "OLED SSD1306\nQR + Amount", "#E9F6F7", "#176B73"),
        "webhook": (1210, 365, 1510, 565, "REST Receiver\nLogs + HTTP 202", "#F0F2F5", "#5B6878"),
    }
    for x1, y1, x2, y2, label, fill, outline in boxes.values():
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline=outline, width=5)
        lines = label.split("\n")
        total_h = sum(draw.textbbox((0, 0), line, font=body_font)[3] for line in lines) + 12 * (len(lines) - 1)
        y = y1 + (y2 - y1 - total_h) / 2
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=body_font)
            draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, fill=outline, font=body_font)
            y += bbox[3] - bbox[1] + 12

    def arrow(start, end, color, label, label_pos):
        draw.line((start, end), fill=color, width=8)
        ex, ey = end
        sx, sy = start
        angle_x = 1 if ex > sx else -1
        draw.polygon([(ex, ey), (ex - 22 * angle_x, ey - 14), (ex - 22 * angle_x, ey + 14)], fill=color)
        draw.text(label_pos, label, fill=color, font=small_font)

    arrow((390, 255), (640, 255), "#2E74B5", "HTTP / Configure / Generate QR", (415, 195))
    arrow((985, 245), (1210, 185), "#08766B", "I2C GPIO 21/22", (1010, 150))
    arrow((985, 350), (1210, 450), "#5B6878", "REST POST JSON", (1015, 395))
    draw.text((65, 530), "ผู้ใช้สร้าง Static/Dynamic QR จาก Dashboard หรือเรียก ESP32 API โดยตรง", fill="#425466", font=small_font)
    draw.text((65, 575), "หมายเหตุ: REST POST ยืนยันเหตุการณ์สร้าง QR ไม่ใช่การยืนยันเงินเข้าบัญชี", fill="#9B1C1C", font=small_font)
    image.save(output)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("PaymentESP | คู่มือการติดตั้งและใช้งาน")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_after = Pt(0)
    run = footer_p.add_run("หน้า ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(footer_p, "PAGE")

    settings = doc.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def add_cover(doc, architecture_path):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(48)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("INSTALLATION & OPERATIONS MANUAL")
    set_run_font(run, size=11, color=GREEN, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("คู่มือการติดตั้งและใช้งาน\nระบบ QR Payment PromptPay ด้วย ESP32")
    set_run_font(run, size=27, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("Visual Studio Code + PlatformIO + Wokwi Simulator + Local Dashboard")
    set_run_font(run, size=13, color=MUTED)

    cover_image = doc.add_paragraph()
    cover_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = cover_image.add_run().add_picture(str(architecture_path), width=Inches(6.1))
    shape._inline.docPr.set("descr", "แผนภาพโครงสร้างการทำงาน PaymentESP")
    shape._inline.docPr.set("title", "แผนภาพโครงสร้างการทำงาน PaymentESP")

    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_before = Pt(18)
    metadata.paragraph_format.space_after = Pt(3)
    run = metadata.add_run("Version 1.0.0 | วันที่ 28 มิถุนายน 2026")
    set_run_font(run, size=10.5, color=MUTED, bold=True)
    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = scope.add_run("เอกสารสำหรับติดตั้ง ทดสอบ และส่งมอบระบบตามขอบเขตใบเสนอราคา")
    set_run_font(run, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


def build_document():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    architecture_path = ASSET_DIR / "paymentesp-architecture.png"
    build_architecture_diagram(architecture_path)

    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "คู่มือการติดตั้งและใช้งาน PaymentESP"
    props.subject = "ESP32 PromptPay QR Payment Installation Manual"
    props.author = "PaymentESP Project Team"
    props.keywords = "ESP32, PromptPay, Wokwi, PlatformIO, OLED, Dashboard"
    props.comments = "Delivery manual version 1.0.0"

    add_cover(doc, architecture_path)

    add_heading(doc, "ข้อมูลเอกสาร", 1)
    add_table(doc, ["รายการ", "รายละเอียด"], [
        ("ชื่อระบบ", "QR Payment PromptPay ด้วย ESP32"),
        ("เวอร์ชัน", "1.0.0"),
        ("วันที่จัดทำ", "28 มิถุนายน 2026"),
        ("สภาพแวดล้อมทดสอบ", "ESP32 DevKit V1, Wokwi Simulator, Visual Studio Code, PlatformIO, Node.js"),
        ("ขอบเขต", "สร้าง Static/Dynamic PromptPay QR, OLED, Serial, REST POST และ Logging"),
    ], [2700, 6660], font_size=10)
    add_callout(doc, "ขอบเขตสำคัญ", "ระบบบันทึกเหตุการณ์สร้าง QR และส่ง REST POST เท่านั้น ไม่ได้ตรวจสอบว่าเงินเข้าบัญชีธนาคารแล้ว", "warning")

    add_heading(doc, "สารบัญ", 1)
    add_list(doc, [
        "ภาพรวมระบบและโครงสร้างไฟล์",
        "สิ่งที่ต้องเตรียมก่อนติดตั้ง",
        "ติดตั้ง Visual Studio Code, PlatformIO และ Wokwi Simulator",
        "Build Firmware และรัน Wokwi",
        "ติดตั้งและใช้งาน Local Dashboard",
        "ตั้งค่า PromptPay ID, Webhook และสร้าง QR",
        "ตรวจ Serial, REST POST และ Logs",
        "ติดตั้งลง ESP32/OLED จริง",
        "ทดสอบรับงานและ Troubleshooting",
        "API Reference และหลักฐานการทดสอบ",
    ])

    doc.add_page_break()
    add_heading(doc, "1. ภาพรวมระบบ", 1)
    add_body(doc, "PaymentESP เป็นระบบสร้าง QR PromptPay บน ESP32 รองรับ QR แบบ Static ที่ไม่ฝังยอดเงิน และ Dynamic ที่ฝังยอดเงินไว้ใน QR ผู้ใช้งานสามารถควบคุมผ่านหน้าเว็บ Local Dashboard หรือเรียก REST API ของ ESP32 โดยตรง")
    add_figure(doc, architecture_path, "ภาพที่ 1 โครงสร้างการทำงานของ PaymentESP", width=6.2)
    add_heading(doc, "1.1 การไหลของข้อมูล", 2)
    add_list(doc, [
        "ผู้ใช้งานกรอก PromptPay ID, Amount และ Reference จาก Dashboard",
        "Dashboard สร้าง QR บนเครื่อง และส่งคำสั่งไป ESP32 เมื่อเปิด Sync",
        "ESP32 สร้าง PromptPay payload, แสดง QR บน OLED และพิมพ์รายละเอียดทาง Serial",
        "ESP32 ส่ง JSON ไปยัง Webhook และบันทึกผล HTTP status ใน Device Logs",
        "Dashboard แสดง Device Logs และ REST POST Events สำหรับตรวจสอบย้อนหลัง",
    ], ordered=True)

    add_heading(doc, "1.2 โครงสร้างชุดส่งมอบ", 2)
    add_table(doc, ["โฟลเดอร์", "เนื้อหา"], [
        ("source/", "Source Code ESP32, Dashboard, Tests และ Config สำหรับ PlatformIO/Wokwi"),
        ("firmware/", "firmware.bin, firmware.elf, bootloader.bin และ partitions.bin"),
        ("documents/", "คู่มือติดตั้ง, Test Report, Requirement Checklist และ Build Summary"),
        ("evidence/", "ภาพ OLED, Serial, Device Logs, REST POST และ event data"),
    ], [2100, 7260])

    add_heading(doc, "2. สิ่งที่ต้องเตรียม", 1)
    add_table(doc, ["รายการ", "ขั้นต่ำ", "หมายเหตุ"], [
        ("ระบบปฏิบัติการ", "Windows 10/11, macOS หรือ Linux", "ต้องติดตั้ง Driver USB เมื่อใช้บอร์ดจริง"),
        ("Visual Studio Code", "เวอร์ชันปัจจุบัน", "ใช้เปิด Source และสั่ง Extension"),
        ("PlatformIO IDE", "Extension", "ใช้ Build, Upload และ Serial Monitor"),
        ("Wokwi Simulator", "Extension", "ใช้ทดสอบ ESP32/OLED โดยไม่ต้องมี Hardware"),
        ("Node.js", "18 หรือใหม่กว่า", "ใช้รัน Local Dashboard"),
        ("ESP32 DevKit V1", "ใช้เมื่อทดสอบ Hardware จริง", "ไม่จำเป็นสำหรับ Wokwi"),
        ("OLED SSD1306", "128x64 I2C, address 0x3C", "ใช้ GPIO 21/22"),
    ], [2100, 2700, 4560], font_size=9.2)
    add_callout(doc, "ก่อนเริ่ม", "แตกไฟล์ ZIP ไปยัง path ที่ไม่มีอักขระพิเศษมากเกินไป และอย่าเปิดโปรเจกต์จากภายใน ZIP โดยตรง", "note")

    add_heading(doc, "3. ติดตั้งเครื่องมือใน Visual Studio Code", 1)
    add_heading(doc, "3.1 เปิด Source Code", 2)
    add_list(doc, [
        "เปิด Visual Studio Code",
        "เลือก File > Open Folder",
        "เลือกโฟลเดอร์ PaymentESP-v1.0.0-20260628/source",
        "กด Trust เมื่อ VS Code ถามเรื่อง Workspace Trust",
    ], ordered=True)
    add_heading(doc, "3.2 ติดตั้ง Extensions", 2)
    add_list(doc, [
        "เปิด Extensions ด้วย Ctrl+Shift+X หรือ Cmd+Shift+X",
        "ค้นหาและติดตั้ง PlatformIO IDE",
        "ค้นหาและติดตั้ง Wokwi Simulator",
        "Reload Window หลังติดตั้งเสร็จ",
    ], ordered=True)
    add_callout(doc, "ตรวจสอบ", "เมื่อ PlatformIO พร้อม จะเห็นไอคอน PlatformIO ที่แถบด้านซ้าย และคำสั่ง PlatformIO: Build ใน Command Palette", "success")

    add_heading(doc, "4. ติดตั้ง Node.js Dependencies", 1)
    add_body(doc, "เปิด Terminal ใน VS Code ที่โฟลเดอร์ source แล้วติดตั้ง dependency ตาม package-lock.json")
    add_code_block(doc, "npm ci")
    add_body(doc, "หากเครื่องไม่มีคำสั่ง npm ให้ติดตั้ง Node.js LTS ก่อน แล้วปิดและเปิด VS Code ใหม่")
    add_callout(doc, "ความปลอดภัย", "ห้ามนำไฟล์ .env หรือ Secret Key จากระบบอื่นใส่ใน Source ที่ส่งต่อ ระบบตามใบเสนอราคาไม่ต้องใช้ Payment Gateway Key", "warning")

    add_heading(doc, "5. Build Firmware ด้วย PlatformIO", 1)
    add_heading(doc, "5.1 Build จาก Command Palette", 2)
    add_list(doc, [
        "กด F1 หรือ Ctrl/Cmd+Shift+P",
        "เลือก PlatformIO: Build",
        "รอจน Terminal แสดง SUCCESS",
    ], ordered=True)
    add_heading(doc, "5.2 Build จาก Terminal", 2)
    add_code_block(doc, "pio run")
    add_body(doc, "ผล Build จะอยู่ที่ .pio/build/esp32dev/firmware.bin และ firmware.elf")
    add_table(doc, ["ค่าจากรอบส่งมอบ", "ผล"], [
        ("Target", "esp32dev / Arduino"),
        ("RAM", "48,448 / 327,680 bytes (14.8%)"),
        ("Flash", "1,016,721 / 1,310,720 bytes (77.6%)"),
        ("Build result", "SUCCESS"),
    ], [3600, 5760], font_size=10)

    add_heading(doc, "6. รัน Wokwi Simulator", 1)
    add_list(doc, [
        "Build firmware ให้ผ่านก่อน",
        "เปิดไฟล์ diagram.json",
        "กด F1 แล้วเลือก Wokwi: Start Simulator",
        "รอ ESP32 boot และเชื่อมต่อ Wokwi-GUEST",
        "เปิด http://localhost:8180 เพื่อตรวจหน้าเว็บของ ESP32",
    ], ordered=True)
    add_heading(doc, "6.1 Port ที่ใช้", 2)
    add_table(doc, ["Port", "หน้าที่", "URL/การใช้งาน"], [
        ("8180", "Wokwi forward ไป ESP32 port 80", "http://localhost:8180"),
        ("4000", "RFC2217 Serial", "Baud 115200"),
        ("3001", "Local Dashboard และ Webhook receiver", "http://localhost:3001"),
    ], [1200, 3600, 4560], font_size=10)
    add_callout(doc, "สถานะพร้อมใช้งาน", "GET http://localhost:8180/api/config ต้องตอบ JSON และมี clockSynced เป็น true เมื่อ NTP สำเร็จ", "success")

    add_heading(doc, "7. รัน Local Dashboard", 1)
    add_code_block(doc, "npm run dashboard")
    add_body(doc, "เปิด Browser ที่ http://localhost:3001 หน้า Dashboard สามารถสร้าง QR บนเครื่องได้ แม้ไม่ได้เปิด Wokwi หากต้องการส่งคำสั่งไป ESP32 ให้เปิด Sync to ESP32 / Wokwi")
    add_heading(doc, "7.1 ตั้งค่าหลัก", 2)
    add_table(doc, ["ช่อง", "ตัวอย่าง", "คำอธิบาย"], [
        ("ESP32 / Wokwi URL", "http://localhost:8180", "ปลายทาง ESP32 ที่ Dashboard จะเรียก"),
        ("PromptPay ID", "0812345678", "เบอร์มือถือ 10 หลัก, เลข 13 หลัก หรือ e-wallet 15 หลัก"),
        ("Webhook URL", "http://host.wokwi.internal:3001/api/webhook", "Wokwi ส่ง event กลับเข้า Dashboard"),
        ("Amount", "15.00", "ใช้กับ Dynamic QR เท่านั้น"),
        ("Reference", "ORDER-0001", "รหัสอ้างอิงสำหรับ Logs"),
    ], [2350, 2800, 4210], font_size=9.2)

    add_heading(doc, "8. บันทึก PromptPay ID และ Webhook", 1)
    add_list(doc, [
        "กรอก PromptPay ID ที่ต้องการ",
        "กรอก Webhook URL",
        "เปิด Sync หากต้องการบันทึกเข้า ESP32/Wokwi",
        "กด Save Settings",
        "ตรวจ Response ว่า saved และ device ไม่มี error",
    ], ordered=True)
    add_body(doc, "Dashboard บันทึกค่าบนเครื่องไว้ที่ local-dashboard/data/config.json ส่วน ESP32 จริงบันทึก PromptPay ID และ Webhook ลง NVS เมื่อเรียก POST /api/config")
    add_callout(doc, "Wokwi", "Wokwi อาจเริ่ม flash session ใหม่หลังปิด Simulator จึงควรเปิด Sync และกด Save Settings เมื่อเริ่มรอบทดสอบใหม่", "note")

    add_heading(doc, "9. สร้าง Static QR และ Dynamic QR", 1)
    add_heading(doc, "9.1 Static QR", 2)
    add_body(doc, "กด Static QR no Amount ระบบจะสร้าง QR ที่ไม่มี Tag ยอดเงิน ผู้สแกนสามารถกรอกยอดในแอปธนาคารเอง")
    add_heading(doc, "9.2 Dynamic QR", 2)
    add_body(doc, "กรอก Amount มากกว่า 0 แล้วกด Dynamic QR with Amount ระบบจะฝังยอดเงินใน QR แอปธนาคารจะแสดงยอดตามที่กำหนดและโดยทั่วไปไม่ให้ผู้จ่ายแก้ไข")
    add_callout(doc, "หลักการ", "Dynamic หมายถึงระบบสร้าง QR ตามยอดที่ระบุ ไม่ได้หมายถึงผู้จ่ายสามารถแก้ยอดในแอปธนาคาร", "note")
    add_heading(doc, "9.3 ทดสอบด้วยแอปธนาคาร", 2)
    add_list(doc, [
        "เปิด QR ที่สร้างบน Dashboard หรือ OLED",
        "สแกนด้วยแอปธนาคาร",
        "ตรวจชื่อผู้รับก่อนดำเนินการทุกครั้ง",
        "ตรวจยอดเงินสำหรับ Dynamic QR",
        "ยกเลิกก่อนยืนยัน หากเป็นการทดสอบที่ไม่ต้องการโอนเงินจริง",
    ], ordered=True)

    add_heading(doc, "10. ตรวจ Serial, REST POST และ Logs", 1)
    add_heading(doc, "10.1 Serial Output", 2)
    add_body(doc, "ทุกครั้งที่ ESP32 สร้าง QR จะพิมพ์ Mode, PromptPay ID, Amount, Reference, เวลา, HTTP status และ Payload ที่ baud 115200")
    add_code_block(doc, "=== QR Payment Event ===\nMode: dynamic\nPromptPay ID: 0066xxxxxxxxx\nAmount: 15.00\nReference: ORDER-0001\nWebhook status: 202\n========================")
    add_heading(doc, "10.2 REST POST", 2)
    add_body(doc, "ESP32 ส่ง event JSON ไปยัง Webhook URL หลังสร้าง QR ค่า HTTP 202 หมายถึง Dashboard receiver รับ event แล้ว")
    add_heading(doc, "10.3 Device Logs", 2)
    add_body(doc, "GET /api/logs แสดงรายการล่าสุดสูงสุด 20 รายการ พร้อม mode, amount, reference, เวลา และ webhookStatus")
    add_callout(doc, "อย่าสับสน", "REST POST 202 และ Logs ยืนยันการสร้าง QR เท่านั้น ไม่ใช่หลักฐานว่าเงินเข้าบัญชีแล้ว", "danger")

    add_heading(doc, "11. API Reference", 1)
    add_table(doc, ["Method", "Endpoint", "หน้าที่", "ตัวอย่าง Parameter"], [
        ("GET", "/api/static", "สร้าง Static QR", "ref=STATIC-TEST"),
        ("GET", "/api/dynamic", "สร้าง Dynamic QR", "amount=15.00&ref=ORDER-1"),
        ("GET", "/api/logs", "ดู Device Logs", "ไม่มี"),
        ("GET", "/api/config", "ดู Config", "ไม่มี"),
        ("POST", "/api/config", "บันทึก PromptPay/Webhook", "promptpay, webhook"),
        ("GET", "/api/qr.svg", "แปลง Payload เป็น SVG", "data=<payload>"),
        ("POST", "Dashboard /api/webhook", "รับ event จาก ESP32", "JSON event"),
        ("GET", "Dashboard /api/webhook/logs", "ดู REST POST Events", "ไม่มี"),
    ], [1100, 2350, 3370, 2540], font_size=8.8)
    add_heading(doc, "11.1 ตัวอย่างคำสั่ง", 2)
    add_code_block(doc, "curl \"http://localhost:8180/api/static?ref=STATIC-TEST\"\ncurl \"http://localhost:8180/api/dynamic?amount=15.00&ref=DYNAMIC-TEST\"\ncurl \"http://localhost:8180/api/logs\"\ncurl \"http://localhost:3001/api/webhook/logs\"")

    add_heading(doc, "12. ติดตั้งลง ESP32 และ OLED จริง", 1)
    add_heading(doc, "12.1 การต่อสาย OLED", 2)
    add_table(doc, ["OLED SSD1306", "ESP32", "หน้าที่"], [
        ("VCC", "3V3", "ไฟเลี้ยง 3.3V"),
        ("GND", "GND", "กราวด์"),
        ("SDA", "GPIO 21", "I2C Data"),
        ("SCL", "GPIO 22", "I2C Clock"),
    ], [2500, 2500, 4360], font_size=10)
    add_callout(doc, "ข้อควรระวัง", "ตรวจสเปกโมดูล OLED ก่อนจ่ายไฟ ห้ามต่อ VCC ผิดขา และควรถอด USB ก่อนเปลี่ยนสาย", "warning")
    add_heading(doc, "12.2 Upload Firmware", 2)
    add_list(doc, [
        "แก้ WIFI_SSID และ WIFI_PASSWORD ใน src/main.cpp ให้ตรงกับสถานที่ใช้งาน",
        "ต่อ ESP32 กับ USB และเลือก Serial Port",
        "รัน pio run -t upload",
        "เปิด Serial Monitor ด้วย pio device monitor -b 115200",
        "ตรวจ IP Address ที่ ESP32 ได้รับ",
    ], ordered=True)
    add_code_block(doc, "pio run -t upload\npio device monitor -b 115200")
    add_heading(doc, "12.3 Webhook สำหรับ Hardware จริง", 2)
    add_body(doc, "เปลี่ยน host.wokwi.internal เป็น IP หรือ Domain ของ Server ที่ ESP32 เข้าถึงได้ เช่น http://192.168.1.10:3001/api/webhook และอนุญาต Firewall port 3001 เฉพาะ Network ที่ต้องใช้")

    add_heading(doc, "13. Automated Tests", 1)
    add_code_block(doc, "g++ -std=c++17 -Iinclude tests/test_promptpay.cpp src/PromptPayQR.cpp -o /tmp/paymentesp-test\n/tmp/paymentesp-test\nnpm run test:dashboard\npio run")
    add_table(doc, ["รายการทดสอบ", "ผลรอบส่งมอบ"], [
        ("PromptPay CRC และ Payload", "PASS"),
        ("Static/Dynamic QR", "PASS"),
        ("PromptPay ID formats", "PASS"),
        ("Dashboard integration", "PASS"),
        ("PlatformIO firmware build", "PASS"),
        ("REST POST / Logging", "PASS"),
    ], [6500, 2860], font_size=10)

    add_heading(doc, "14. Acceptance Test ก่อนใช้งาน", 1)
    add_list(doc, [
        "Build firmware แสดง SUCCESS",
        "Wokwi หรือ ESP32 boot และเชื่อม WiFi สำเร็จ",
        "/api/config ตอบ JSON",
        "Dashboard เปิดที่ localhost:3001",
        "Save PromptPay ID และ Webhook สำเร็จ",
        "Static QR สแกนได้และไม่มียอดฝัง",
        "Dynamic QR สแกนได้และแสดงยอดตรงกับ Dashboard",
        "Serial แสดง reference และ amount ถูกต้อง",
        "Device Logs มีรายการล่าสุด",
        "REST POST Events มี reference เดียวกันและ HTTP 202",
        "ชื่อผู้รับในแอปธนาคารถูกต้อง",
    ], ordered=True)
    add_callout(doc, "เกณฑ์ผ่าน", "ทุกข้อด้านบนต้องผ่านก่อนนำไปสาธิตหรือส่งมอบหน้างาน", "success")

    add_heading(doc, "15. Troubleshooting", 1)
    add_table(doc, ["อาการ", "สาเหตุที่พบบ่อย", "วิธีแก้"], [
        ("Device request timed out", "Wokwi ยังไม่ boot หรือ port 8180 ค้าง", "Build ใหม่, Start Simulator, Reload Window"),
        ("localhost:8180 เปิดไม่ได้", "Simulator ไม่ได้รัน", "เปิด diagram.json แล้วสั่ง Wokwi: Start Simulator"),
        ("Dashboard เปิดไม่ได้", "Node server ไม่ได้รันหรือ port ถูกใช้", "รัน npm run dashboard หรือตั้ง DASHBOARD_PORT ใหม่"),
        ("QR ไม่มี Amount", "เลือก Static QR", "เลือก Dynamic QR และกรอก Amount > 0"),
        ("REST POST ว่าง", "Webhook URL ผิดหรือ Dashboard ไม่ได้รัน", "ใช้ host.wokwi.internal:3001 สำหรับ Wokwi"),
        ("OLED ไม่แสดง", "สาย SDA/SCL หรือ address ผิด", "ตรวจ GPIO 21/22 และ address 0x3C"),
        ("PromptPay ID กลับค่าเดิม", "Wokwi เริ่ม flash session ใหม่", "เปิด Sync แล้วกด Save Settings อีกครั้ง"),
        ("Serial ไม่มีข้อความ", "Monitor/port ไม่ถูกต้อง", "เลือก port และ baud 115200 แล้ว Restart Simulator"),
    ], [2500, 3480, 3380], font_size=8.8)
    add_heading(doc, "15.1 ตรวจสุขภาพระบบแบบเร็ว", 2)
    add_code_block(doc, "curl http://localhost:8180/api/config\ncurl http://localhost:8180/api/logs\ncurl http://localhost:3001/api/webhook/logs")

    add_heading(doc, "16. Security และขอบเขตงาน", 1)
    add_list(doc, [
        "อย่า Commit หรือส่งไฟล์ .env, Secret Key และข้อมูลทดสอบส่วนบุคคล",
        "ตรวจชื่อผู้รับในแอปธนาคารทุกครั้งก่อนชำระ",
        "จำกัดการเข้าถึง Dashboard และ Webhook เมื่อใช้งานบน Network จริง",
        "ตั้ง Reference ไม่ซ้ำหากระบบภายนอกต้องติดตามรายการ",
        "สำรอง Source Code และ Config ก่อนแก้ Firmware",
    ])
    add_heading(doc, "16.1 สิ่งที่ไม่รวมตามใบเสนอราคา", 2)
    add_list(doc, [
        "Hardware ESP32/OLED จริง",
        "Hosting หรือ Production Server",
        "ระบบตรวจสอบเงินเข้าบัญชีอัตโนมัติ",
        "Bank API หรือ Payment Gateway production",
    ])
    add_callout(doc, "งานระยะถัดไป", "หากต้องการให้ตู้รับทราบผลชำระเงินจริง ต้องเชื่อม Bank API หรือ Payment Gateway ที่ส่ง Webhook การชำระสำเร็จ", "note")

    doc.add_page_break()
    add_heading(doc, "ภาคผนวก A: หลักฐานการทดสอบ", 1)
    add_body(doc, "หลักฐานต่อไปนี้ใช้ reference EVIDENCE-001 ยอด 15.00 บาท และบันทึกจากรอบทดสอบ Wokwi/Local Dashboard")
    figures = [
        ("01-oled-display.png", "ภาพที่ A-1 OLED แสดง Dynamic QR, Amount และ Reference"),
        ("02-serial-output.png", "ภาพที่ A-2 Serial Output ของเหตุการณ์สร้าง QR"),
        ("03-device-logs.png", "ภาพที่ A-3 Device Logs พร้อม HTTP POST status 202"),
        ("04-rest-post-events.png", "ภาพที่ A-4 REST POST Events ที่ Dashboard รับสำเร็จ"),
    ]
    for filename, caption in figures:
        add_figure(doc, EVIDENCE_DIR / filename, caption, width=6.15)
    add_callout(doc, "สรุปผล", "OLED, Serial, Device Logs และ REST POST มี Amount/Reference ตรงกัน ผลทดสอบ PASS", "success")

    add_heading(doc, "ภาคผนวก B: ข้อมูลติดต่อและการส่งมอบ", 1)
    add_body(doc, "เริ่มต้นใช้งานจากไฟล์ README.md ในชุดส่งมอบ และเก็บ SHA256SUMS.txt ไว้ใช้ตรวจความถูกต้องของไฟล์")
    add_table(doc, ["เอกสาร", "ตำแหน่ง"], [
        ("คู่มือนี้", "documents/PaymentESP-Detailed-Installation-Manual-TH.docx"),
        ("Test Report", "documents/test-report.md"),
        ("Requirement Checklist", "documents/requirements-checklist.md"),
        ("Source Code", "source/"),
        ("Firmware", "firmware/"),
        ("Evidence", "evidence/"),
    ], [3000, 6360], font_size=10)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
