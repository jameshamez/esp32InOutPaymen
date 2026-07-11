#!/usr/bin/env python3
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_installation_docx import (
    BLACK,
    DARK_BLUE,
    GREEN,
    MUTED,
    NAVY,
    add_body,
    add_callout,
    add_code_block,
    add_figure,
    add_heading,
    add_list,
    add_table,
    configure_document,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
VERSION = "1.1.0-Hardware / Rev.3"
PACKAGE_NAME = "PaymentESP-v1.1.0-Hardware-20260628"
OUTPUT = ROOT / "deliverables" / PACKAGE_NAME / "documents" / "PaymentESP-Hardware-Installation-Manual-TH.docx"
ASSET_DIR = ROOT / "docs" / "word-assets"


def build_hardware_diagram(output):
    width, height = 1600, 700
    image = Image.new("RGB", (width, height), "#F5F7FA")
    draw = ImageDraw.Draw(image)
    font_path = "/System/Library/AssetsV2/com_apple_MobileAsset_Font8/cf0dc8d3b09f9ba379660e591e82566e2b557949.asset/AssetData/Sarabun.ttc"
    title_font = ImageFont.truetype(font_path, 42)
    body_font = ImageFont.truetype(font_path, 30)
    small_font = ImageFont.truetype(font_path, 23)
    draw.text((55, 32), "การติดตั้ง PaymentESP บน ESP32 จริง", fill="#17324D", font=title_font)

    boxes = [
        (60, 180, 360, 420, "โทรศัพท์ / PC\nตั้งค่าครั้งแรก\n192.168.4.1", "#E8EEF5", "#2E74B5"),
        (540, 145, 900, 455, "ESP32 DevKit V1\nPaymentESP Firmware\nNVS + Web API", "#EAF8F4", "#08766B"),
        (1120, 80, 1500, 280, "OLED SSD1306\nQR + Amount\nGPIO 21 / 22", "#E9F6F7", "#176B73"),
        (1120, 370, 1500, 570, "WiFi Router / LAN\npaymentesp.local\nDashboard / REST", "#F0F2F5", "#5B6878"),
    ]
    for x1, y1, x2, y2, label, fill, outline in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, fill=fill, outline=outline, width=5)
        lines = label.split("\n")
        total_h = sum(draw.textbbox((0, 0), line, font=body_font)[3] for line in lines) + 8 * (len(lines) - 1)
        y = y1 + (y2 - y1 - total_h) / 2
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=body_font)
            draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, fill=outline, font=body_font)
            y += bbox[3] - bbox[1] + 8

    def arrow(start, end, color, label, label_pos):
        draw.line((start, end), fill=color, width=8)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 22, ey - 14), (ex - 22, ey + 14)], fill=color)
        draw.text(label_pos, label, fill=color, font=small_font)

    arrow((360, 270), (540, 270), "#2E74B5", "Setup AP / HTTP", (370, 215))
    arrow((900, 235), (1120, 180), "#08766B", "I2C", (970, 155))
    arrow((900, 365), (1120, 455), "#5B6878", "WiFi / REST", (950, 405))
    draw.text((55, 605), "ครั้งแรก: เชื่อม PaymentESP-Setup แล้วกรอก WiFi/PromptPay ที่ 192.168.4.1/setup", fill="#425466", font=small_font)
    draw.text((55, 645), "ใช้งานปกติ: เปิด paymentesp.local หรือ IP จาก OLED โดยอยู่ใน WiFi เดียวกับ ESP32", fill="#08766B", font=small_font)
    image.save(output)


def add_cover(doc, architecture_path):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(45)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    set_run_font(kicker.add_run("HARDWARE INSTALLATION & HANDOVER MANUAL"), size=11, color=GREEN, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    set_run_font(title.add_run("คู่มือติดตั้งและส่งมอบ\nPaymentESP บน ESP32 จริง"), size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    set_run_font(subtitle.add_run("Source Code Setup + ESP32 DevKit V1 + OLED SSD1306 + WiFi Setup Portal"), size=13, color=MUTED)

    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = image_p.add_run().add_picture(str(architecture_path), width=Inches(6.15))
    shape._inline.docPr.set("descr", "แผนภาพการติดตั้ง PaymentESP บน ESP32 และ OLED จริง")
    shape._inline.docPr.set("title", "PaymentESP Hardware Architecture")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(16)
    set_run_font(meta.add_run(f"Version {VERSION} | 28 มิถุนายน 2026"), size=10.5, color=MUTED, bold=True)
    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(scope.add_run("ฉบับละเอียดสำหรับเตรียม Source Code, Upload เครื่องจริง, ตั้งค่าหน้างาน และตรวจรับก่อนส่งลูกค้า"), size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


def build_document():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    architecture_path = ASSET_DIR / "paymentesp-hardware-architecture.png"
    build_hardware_diagram(architecture_path)

    doc = Document()
    configure_document(doc)
    doc.sections[0].header.paragraphs[0].clear()
    header = doc.sections[0].header.paragraphs[0]
    set_run_font(header.add_run("PaymentESP | คู่มือติดตั้ง ESP32 Hardware"), size=8.5, color=MUTED, bold=True)
    props = doc.core_properties
    props.title = "คู่มือติดตั้ง PaymentESP บน ESP32 Hardware"
    props.subject = "ESP32 PromptPay QR hardware installation and handover"
    props.author = "PaymentESP Project Team"
    props.keywords = "ESP32, PromptPay, Hardware, OLED, PlatformIO, WiFi Setup"
    props.comments = f"Hardware delivery manual version {VERSION}"

    add_cover(doc, architecture_path)

    add_heading(doc, "ข้อมูลเอกสาร", 1)
    add_table(doc, ["รายการ", "รายละเอียด"], [
        ("ชื่อระบบ", "QR Payment PromptPay ด้วย ESP32"),
        ("เวอร์ชัน", VERSION),
        ("เป้าหมายติดตั้ง", "ESP32 DevKit V1 และ OLED SSD1306 128x64 I2C"),
        ("Source Code root", "โฟลเดอร์ source/ ที่มีไฟล์ platformio.ini"),
        ("การตั้งค่าครั้งแรก", "WiFi AP PaymentESP-Setup และหน้า 192.168.4.1/setup"),
        ("การเข้าใช้งาน", "http://paymentesp.local หรือ IP ที่แสดงบน OLED/Serial"),
        ("สถานะ Upload รอบจัดทำ", "รอต่อบอร์ด ESP32 จริง; Firmware Build ผ่านแล้ว"),
    ], [2700, 6660], font_size=10)
    add_callout(doc, "สำคัญ", "Firmware และชุดติดตั้งพร้อมแล้ว แต่ต้อง Upload และทำ Hardware Acceptance กับบอร์ด/OLED ตัวที่จะส่งลูกค้าก่อนปิดกล่อง", "warning")

    add_heading(doc, "สารบัญ", 1)
    add_list(doc, [
        "ภาพรวมการติดตั้ง Hardware และรายการอุปกรณ์",
        "การต่อสาย ESP32, OLED และ Payment Pulse",
        "เตรียมคอมพิวเตอร์ ติดตั้ง Visual Studio Code/PlatformIO และ Setup Source Code",
        "โครงสร้าง Source, Environment, Dependencies และการเก็บข้อมูลลับ",
        "อธิบาย Code: Boot, NVS, PromptPay TLV/CRC, API, OLED, Logs และ Dashboard",
        "Build และ Upload Firmware ลงบอร์ดจริง",
        "ตั้งค่า WiFi และ PromptPay ผ่าน Setup AP",
        "เข้าใช้งานผ่าน paymentesp.local และ Local Dashboard",
        "สร้าง Static/Dynamic QR และตรวจ Output",
        "Troubleshooting และวิธีกู้คืน WiFi",
        "Hardware Acceptance ก่อนส่งลูกค้า",
        "ไฟล์ส่งมอบและหลักฐานการทดสอบ",
    ])

    doc.add_page_break()
    add_heading(doc, "1. ภาพรวม Hardware", 1)
    add_body(doc, "PaymentESP ทำงานบน ESP32 จริงโดยเก็บ WiFi, PromptPay ID และ Webhook URL ใน NVS เมื่อยังไม่มี WiFi หรือเชื่อมต่อไม่ได้ บอร์ดจะเปิด Setup AP เพื่อให้ตั้งค่าจากโทรศัพท์หรือคอมพิวเตอร์โดยไม่ต้องแก้ Source Code")
    add_figure(doc, architecture_path, "ภาพที่ 1 การติดตั้ง PaymentESP บน ESP32 จริง", width=6.2)
    add_heading(doc, "1.1 ลำดับการทำงาน", 2)
    add_list(doc, [
        "ช่าง Upload firmware environment esp32dev ลงบอร์ด",
        "ESP32 เปิด PaymentESP-Setup หากยังไม่มี WiFi ที่ใช้งานได้",
        "ผู้ติดตั้งบันทึก WiFi, PromptPay ID และ Webhook ผ่านหน้า setup",
        "ESP32 รีสตาร์ต เชื่อม Router และแสดง IP บน OLED/Serial",
        "ลูกค้าเปิด paymentesp.local หรือ IP เพื่อสร้าง QR และตรวจ Logs",
    ], ordered=True)
    doc.add_page_break()
    add_heading(doc, "1.2 อุปกรณ์ที่ต้องใช้", 2)
    add_table(doc, ["รายการ", "สเปก", "จำนวน"], [
        ("ESP32", "ESP32 DevKit V1, Flash 4 MB", "1"),
        ("OLED", "SSD1306 128x64 I2C, address 0x3C", "1"),
        ("สาย USB", "สาย Data ที่รองรับ Upload", "1"),
        ("สาย Jumper", "Female-Female หรือให้ตรงกับโมดูล", "4 เส้น"),
        ("WiFi", "2.4 GHz พร้อม Internet สำหรับ NTP/Webhook", "1 Network"),
    ], [2500, 5160, 1700], font_size=9.5)

    add_heading(doc, "2. การต่อสาย", 1)
    add_table(doc, ["OLED SSD1306", "ESP32", "หน้าที่"], [
        ("VCC", "3V3", "ไฟเลี้ยง 3.3V"),
        ("GND", "GND", "กราวด์"),
        ("SDA", "GPIO 21", "I2C Data"),
        ("SCL", "GPIO 22", "I2C Clock"),
    ], [2500, 2500, 4360], font_size=10)
    add_callout(doc, "ความปลอดภัย", "ถอดสาย USB ก่อนเปลี่ยนสาย ตรวจตำแหน่ง VCC/GND จากสกรีนของโมดูล และอย่าจ่าย 5V หากโมดูลระบุให้ใช้ 3.3V เท่านั้น", "warning")
    add_heading(doc, "2.1 Output Pulse", 2)
    add_body(doc, "GPIO 26 ถูกกำหนดเป็น Payment Pulse Output ระยะ 500 ms สำหรับระบบภายนอกที่เรียก POST /api/payment ด้วยสถานะ paid ขานี้ไม่ใช่สัญญาณยืนยันเงินจากธนาคารด้วยตัวเอง")

    doc.add_page_break()
    add_heading(doc, "3. เตรียมคอมพิวเตอร์และ Setup Source Code", 1)
    add_body(doc, "ขั้นตอนนี้ใช้สำหรับเครื่องช่างที่ Build และ Upload Firmware ให้ ESP32 หากส่งเฉพาะไฟล์ firmware-merged.bin สามารถข้ามไปหัวข้อ 4.6 ได้ แต่การเก็บ Source Code ไว้ช่วยให้แก้ไขและ Build ซ้ำในอนาคต")
    add_heading(doc, "3.1 Software ที่ต้องติดตั้ง", 2)
    add_table(doc, ["Software", "เวอร์ชันแนะนำ", "ใช้ทำอะไร"], [
        ("Visual Studio Code", "รุ่นปัจจุบัน", "เปิดและแก้ไข Source Code"),
        ("PlatformIO IDE", "Extension ใน VS Code", "Build, Upload และ Serial Monitor"),
        ("USB Driver", "CP210x หรือ CH340", "ทำให้ระบบเห็น Serial Port ของ ESP32"),
        ("Node.js", "18 ขึ้นไป", "รัน Local Dashboard และ tests"),
        ("Git", "ไม่บังคับ", "สำรองเวอร์ชัน Source Code"),
    ], [2400, 2400, 4560], font_size=9.3)
    add_callout(doc, "สาย USB", "ต้องใช้สาย USB Data สายชาร์จอย่างเดียวจะจ่ายไฟได้แต่ไม่แสดง Serial Port และ Upload ไม่ได้", "warning")

    add_heading(doc, "3.2 แตกไฟล์และเปิดโฟลเดอร์ที่ถูกต้อง", 2)
    add_list(doc, [
        "แตกไฟล์ PaymentESP-v1.1.0-Hardware-20260628.zip ลงโฟลเดอร์ที่เขียนไฟล์ได้",
        "เปิด Visual Studio Code แล้วเลือก File > Open Folder",
        "เลือกโฟลเดอร์ source โดยตรง ไม่เลือกโฟลเดอร์ ZIP ชั้นนอก",
        "ตรวจว่าแถบ Explorer เห็น platformio.ini, src/, include/, tests/ และ local-dashboard/",
        "รอ PlatformIO แสดง Project Tasks และดาวน์โหลด toolchain/libraries ครั้งแรกจนเสร็จ",
    ], ordered=True)
    add_code_block(doc, "PaymentESP-v1.1.0-Hardware-20260628/\n└── source/                 <- Open Folder ที่นี่\n    ├── platformio.ini\n    ├── src/\n    ├── include/\n    ├── tests/\n    └── local-dashboard/")
    add_callout(doc, "ตรวจจุดเปิดโครงการ", "ถ้า PlatformIO ไม่พบ Project ให้ตรวจว่าไฟล์ platformio.ini อยู่ระดับบนสุดของโฟลเดอร์ที่ VS Code เปิด", "note")

    add_heading(doc, "3.3 โครงสร้าง Source Code", 2)
    add_table(doc, ["ไฟล์/โฟลเดอร์", "หน้าที่", "แก้เมื่อใด"], [
        ("platformio.ini", "กำหนด Board, Framework, Libraries และ Environment", "เปลี่ยน Board/Upload speed หรือ dependency"),
        ("src/main.cpp", "WiFi, Setup AP, Web API, OLED, NVS, Serial และ Pulse", "แก้พฤติกรรมหลักหรือ GPIO"),
        ("src/PromptPayQR.cpp", "สร้าง PromptPay payload และ CRC", "แก้ logic QR โดยต้องรัน unit test"),
        ("include/PromptPayQR.h", "โครงสร้างข้อมูลและ function declarations", "เมื่อเปลี่ยน interface ของ QR module"),
        ("local-dashboard/", "หน้าเว็บบนคอมพิวเตอร์และ REST receiver", "แก้ UI/Proxy/Logging ฝั่ง PC"),
        ("tests/", "Automated tests", "รันก่อนส่งมอบทุกครั้ง"),
        ("diagram.json / wokwi.toml", "Wokwi Simulator เท่านั้น", "ไม่ใช้ Upload เครื่องจริง"),
    ], [2600, 4200, 2560], font_size=8.5)

    add_heading(doc, "3.4 PlatformIO Environments", 2)
    add_table(doc, ["Environment", "คำสั่ง", "ใช้สำหรับ"], [
        ("esp32dev", "pio run -e esp32dev", "Firmware สำหรับ ESP32 เครื่องจริง; เป็น default environment"),
        ("wokwi", "pio run -e wokwi", "Simulator; มี flag PAYMENTESP_WOKWI"),
    ], [2100, 3300, 3960], font_size=9.2)
    add_callout(doc, "สำคัญ", "เครื่องที่จะส่งลูกค้าต้อง Build/Upload ด้วย esp32dev เท่านั้น เพื่อไม่ให้ Wokwi-GUEST และ host.wokwi.internal ติดไปกับ Firmware", "warning")

    add_heading(doc, "3.5 ติดตั้ง Dependencies", 2)
    add_body(doc, "PlatformIO จะอ่าน lib_deps จาก platformio.ini และติดตั้ง QRCode, Adafruit SSD1306 และ Adafruit GFX โดยอัตโนมัติ หากต้องการสั่งติดตั้งล่วงหน้าให้เปิด Terminal ใน source/ แล้วรันคำสั่งต่อไปนี้")
    add_code_block(doc, "pio pkg install -e esp32dev")
    add_table(doc, ["Library", "Version ในโครงการ", "หน้าที่"], [
        ("ricmoo/QRCode", "0.0.1", "สร้าง QR matrix สำหรับ OLED"),
        ("Adafruit SSD1306", "^2.5.15", "ควบคุมจอ SSD1306"),
        ("Adafruit GFX Library", "^1.12.3", "วาดข้อความและกราฟิก"),
    ], [3100, 2200, 4060], font_size=9.2)

    add_heading(doc, "3.6 ตรวจ Toolchain", 2)
    add_code_block(doc, "pio --version\npio project config\npio run -e esp32dev")
    add_body(doc, "ถ้าคำสั่ง pio ไม่พบ ให้ใช้ PlatformIO sidebar > Project Tasks > esp32dev > General > Build หรือเปิด Terminal จาก PlatformIO extension")

    add_heading(doc, "3.7 ข้อมูลที่ไม่ควร Hardcode ใน Source", 2)
    add_list(doc, [
        "WiFi SSID และ Password ให้กรอกผ่านหน้า Setup Portal",
        "PromptPay ID จริงให้บันทึกผ่านหน้า Setup Portal หรือ /api/config",
        "Webhook URL ให้ใช้ IP/Domain ที่ ESP32 เข้าถึงได้",
        "ห้ามใส่ Secret Key, ไฟล์ .env หรือ WiFi Password ลง Git/ZIP ส่งมอบ",
        "ค่าเริ่มต้น 0812345678 เป็น Demo เท่านั้น ต้องเปลี่ยนก่อนสแกนจริง",
    ])
    add_callout(doc, "NVS", "ค่าหน้างานถูกเก็บใน Preferences namespace promptpay บน Flash ของ ESP32 ไม่ได้เขียนกลับเข้า Source Code", "note")

    add_heading(doc, "3.8 ติดตั้ง USB Driver และตรวจ Serial Port", 2)
    add_table(doc, ["ระบบ", "ตัวอย่าง Port", "วิธีตรวจ"], [
        ("Windows", "COM3", "Device Manager > Ports; ติดตั้ง CP210x/CH340 หากไม่พบ"),
        ("macOS", "/dev/cu.usbserial-*", "รัน pio device list และอนุญาต Driver ใน Privacy & Security"),
        ("Linux", "/dev/ttyUSB0", "รัน pio device list; เพิ่มผู้ใช้เข้า dialout แล้ว login ใหม่"),
    ], [2100, 3300, 3960], font_size=9.5)

    add_heading(doc, "3.9 อธิบาย Source Code", 1)
    add_body(doc, "ส่วนนี้อธิบายโค้ดตามไฟล์จริงในชุดส่งมอบ เพื่อให้ผู้ดูแลระบบสามารถติดตามลำดับการทำงาน แก้ไข และทดสอบซ้ำได้โดยไม่ต้องอ่านทุกไฟล์ตั้งแต่ต้น")

    add_heading(doc, "3.9.1 แบ่งโค้ดออกเป็น 4 ชั้น", 2)
    add_table(doc, ["ชั้น", "ไฟล์หลัก", "ความรับผิดชอบ"], [
        ("Data contract", "include/PromptPayQR.h", "ประกาศ PromptPayConfig, PromptPayPayload, QrEvent และฟังก์ชันที่ใช้ข้ามโมดูล"),
        ("QR core", "src/PromptPayQR.cpp", "Normalize PromptPay ID, สร้าง TLV payload, CRC16 และ JSON โดยไม่ผูกกับ Hardware"),
        ("Device runtime", "src/main.cpp", "WiFi, Setup AP, NVS, Web API, OLED, Serial, Logs, Webhook และ GPIO pulse"),
        ("PC dashboard", "local-dashboard/server.js", "หน้าเว็บบนคอมพิวเตอร์, Local QR, Proxy, Webhook receiver และไฟล์ Log"),
        ("Tests", "tests/", "ตรวจ PromptPay core, Dashboard และ Omise adapter แบบอัตโนมัติ"),
    ], [1900, 2700, 4760], font_size=9.0)
    add_callout(doc, "หลักการสำคัญ", "PromptPayQR.cpp เป็น pure logic จึงทดสอบบนคอมพิวเตอร์ได้ ส่วน main.cpp เป็น Hardware adapter สำหรับ ESP32", "note")

    add_heading(doc, "3.9.2 โครงสร้างข้อมูลหลัก", 2)
    add_table(doc, ["ชนิดข้อมูล", "ฟิลด์สำคัญ", "ใช้เมื่อใด"], [
        ("PromptPayConfig", "promptPayId, merchantName, city", "ค่าที่ใช้เป็นต้นทางในการสร้าง QR"),
        ("PromptPayPayload", "payload, normalizedTarget, targetType", "ผลจากการสร้าง EMV PromptPay payload"),
        ("QrEvent", "mode, amount, reference, time, webhookStatus", "ข้อมูลเหตุการณ์ที่ส่ง Serial, REST, Logs และ OLED"),
    ], [2100, 3500, 3760], font_size=9.2)

    add_heading(doc, "3.9.3 ลำดับ Boot ใน setup() และ loop()", 2)
    add_code_block(doc, "loadPersistedConfig();\nsetupDisplay();\nconnectWifi();\nif (WiFi.status() == WL_CONNECTED) syncClock();\n// register HTTP routes, then server.begin();")
    add_list(doc, [
        "Serial เริ่มที่ 115200 baud เพื่อให้ช่างอ่าน Boot log ได้",
        "loadPersistedConfig() อ่าน WiFi, PromptPay ID และ Webhook จาก NVS",
        "GPIO 26 ถูกตั้งเป็น OUTPUT และเริ่มต้น LOW",
        "setupDisplay() เปิด I2C GPIO 21/22 และจอ SSD1306 address 0x3C",
        "connectWifi() ลองต่อ WiFi 20 วินาที ถ้าไม่สำเร็จจะเปิด PaymentESP-Setup",
        "เมื่อมี Internet ระบบ sync เวลา NTP แล้วเปิด HTTP server port 80",
        "loop() เรียก server.handleClient() และปิด Payment Pulse เมื่อครบ 500 ms",
    ], ordered=True)

    add_heading(doc, "3.9.4 การเก็บค่าหน้างานใน NVS", 2)
    add_body(doc, "Preferences ใช้ namespace promptpay และ key id, webhook, wifiSsid, wifiPass ค่าเหล่านี้อยู่ใน Flash ของบอร์ดและยังคงอยู่หลังปิดเปิดเครื่อง")
    add_code_block(doc, "preferences.begin(PREFERENCES_NAMESPACE, false);\npreferences.putString(PREFERENCES_ID_KEY, qrConfig.promptPayId.c_str());\npreferences.putString(PREFERENCES_WEBHOOK_KEY, webhookUrl.c_str());")
    add_table(doc, ["ฟังก์ชัน", "หน้าที่", "ข้อควรระวัง"], [
        ("loadPersistedConfig()", "อ่านค่าและ validate ก่อนนำไปใช้", "ค่าผิดรูปแบบจะ fallback เป็นค่าเริ่มต้น"),
        ("handleSetup()", "บันทึก WiFi/PromptPay/Webhook แล้ว restart", "SSID เดิมและ Password ว่างจะคงรหัสเดิม; SSID ใหม่ต้องกรอกรหัส"),
        ("handleConfig()", "GET ดูสถานะ; POST บันทึก PromptPay/Webhook", "API ไม่ส่ง WiFi Password กลับออกมา"),
    ], [2400, 3600, 3360], font_size=9.0)

    doc.add_page_break()
    add_heading(doc, "3.9.5 การสร้าง PromptPay Payload", 2)
    add_body(doc, "buildPromptPayPayload() ใน src/PromptPayQR.cpp แปลงข้อมูลเป็น EMV TLV โดย tlv() ต่อ Tag + ความยาว 2 หลัก + ค่า แล้วคำนวณ CRC16-CCITT-FALSE ปิดท้าย")
    add_table(doc, ["Tag", "ค่า/ความหมาย", "พฤติกรรม"], [
        ("00", "Payload Format 01", "ระบุรูปแบบ EMV QR"),
        ("01", "11 Static / 12 Dynamic", "Dynamic ใช้เมื่อระบบกำหนดยอด"),
        ("29", "AID + PromptPay target", "มือถือถูกแปลงเป็น 0066xxxxxxxxx; ID/e-wallet ใช้ sub-tag ต่างกัน"),
        ("58 / 53", "TH / 764", "ประเทศและสกุลเงินบาท"),
        ("54", "Amount ทศนิยม 2 ตำแหน่ง", "ใส่เมื่อมี amount; Dynamic ต้องมากกว่า 0"),
        ("62.05", "Reference", "ตัดไม่เกิน 25 ตัวอักษร"),
        ("63", "CRC16", "ใช้ตรวจว่า payload ไม่เสียหาย"),
    ], [1300, 4000, 4060], font_size=8.8)
    add_code_block(doc, "payload += tlv(\"01\", dynamicQr ? \"12\" : \"11\");\nif (!amountText.empty()) payload += tlv(\"54\", amountText);\npayload += \"6304\";\nsnprintf(crc, sizeof(crc), \"%04X\", crc16CcittFalse(payload));\npayload += crc;")
    add_callout(doc, "Static กับ Dynamic", "Static QR ไม่มี Tag 54 ผู้จ่ายกรอกยอดเอง ส่วน Dynamic QR มี Tag 54 แอปธนาคารจะแสดงยอดที่ระบบกำหนด", "note")

    add_heading(doc, "3.9.6 เส้นทางสร้าง QR Event", 2)
    add_code_block(doc, "buildPromptPayPayload()\n  -> setEventTime()\n  -> postEvent()\n  -> saveLog()\n  -> printEventToSerial()\n  -> renderEventToOled()")
    add_body(doc, "createEvent() เป็นจุดรวมของทุก Output หลังสร้าง payload ระบบบันทึกเวลา ส่ง JSON ไป Webhook เก็บ Log สูงสุด 20 รายการ พิมพ์ Serial และวาด QR บน OLED จาก QrEvent ชุดเดียวกัน จึงทำให้ amount/reference ตรงกันทุกช่องทาง")
    add_callout(doc, "Webhook status", "HTTP 2xx ใน QrEvent หมายถึงปลายทางรับ Event การสร้าง QR แล้ว ไม่ได้ยืนยันว่าเงินเข้าบัญชี", "warning")

    add_heading(doc, "3.9.7 HTTP Route และ Handler", 2)
    add_table(doc, ["Route", "Handler", "สิ่งที่โค้ดทำ"], [
        ("GET /", "handleRoot()", "ส่งหน้าเว็บที่ฝังอยู่ใน Firmware"),
        ("GET/POST /setup", "handleSetup()", "ตั้ง WiFi, PromptPay และ Webhook แล้ว restart"),
        ("GET /api/static", "handleStaticQr()", "สร้าง Event แบบไม่มี amount"),
        ("GET /api/dynamic", "handleDynamicQr()", "validate amount > 0 แล้วสร้าง Dynamic QR"),
        ("GET /api/logs", "handleLogs()", "แปลง QrEvent ใน RAM เป็น JSON"),
        ("GET/POST /api/config", "handleConfig()", "ดูสถานะหรือบันทึก PromptPay/Webhook ลง NVS"),
        ("POST /api/payment", "handlePaymentConfirmation()", "รับ paid callback ที่เชื่อถือได้ แสดง PAID และ Pulse GPIO 26"),
    ], [2300, 2900, 4160], font_size=8.5)

    add_heading(doc, "3.9.8 OLED, QR Matrix และ Payment Pulse", 2)
    add_list(doc, [
        "buildQrCode() ทดลอง QR version 6 ถึง 15 และใช้ ECC_LOW เพื่อให้ payload ยาวพอดีกับจอ",
        "renderEventToOled() จำกัด QR ไม่เกิน 56 pixel และแสดง mode, amount, reference ด้านขวา",
        "renderNetworkStatus() แสดง SETUP/READY, IP และ paymentesp.local ตามสถานะเครือข่าย",
        "handlePaymentConfirmation() ป้องกัน paymentId ซ้ำ ก่อนยก GPIO 26 HIGH 500 ms",
    ])
    add_callout(doc, "ความปลอดภัยของ /api/payment", "Endpoint นี้ต้องรับข้อมูลจาก Payment Gateway/Bank backend ที่ตรวจสอบลายเซ็นแล้ว ห้ามถือ POST ที่ผู้ใช้ส่งเองเป็นหลักฐานการชำระเงิน", "warning")

    doc.add_page_break()
    add_heading(doc, "3.9.9 Local Dashboard ทำงานอย่างไร", 2)
    add_body(doc, "local-dashboard/server.js เป็นบริการ Node.js บนคอมพิวเตอร์ ไม่ได้ทำงานบน ESP32 และไม่จำเป็นต่อการสร้าง QR จากหน้าเว็บในตัวบอร์ด")
    add_table(doc, ["ส่วน", "หน้าที่", "ข้อมูลที่เก็บ"], [
        ("Dashboard config", "เก็บ PromptPay ID/Webhook ฝั่ง PC", "local-dashboard/data/config.json"),
        ("Local QR", "สร้าง payload/SVG บน PC เมื่อ ESP32 ยังไม่ตอบ", "ไม่เขียน NVS จนกด Sync to hardware"),
        ("Proxy", "ส่ง request จาก Browser ไป ESP32", "จำกัด target เพื่อป้องกัน request ออกนอกวงที่อนุญาต"),
        ("Webhook receiver", "รับ Event สร้าง QR และแสดง Evidence", "event.json/HTML และ Log ฝั่ง PC"),
        ("Payment adapter", "เชื่อม Gateway เมื่อกำหนด Secret ผ่าน environment", "Secret ต้องไม่อยู่ใน Source หรือ ZIP"),
    ], [2200, 3700, 3460], font_size=8.8)

    add_heading(doc, "3.9.10 แก้ไขตรงไหนเมื่อเปลี่ยน Requirement", 2)
    add_table(doc, ["ต้องการเปลี่ยน", "ไฟล์/ฟังก์ชัน", "ต้องทดสอบ"], [
        ("รูปแบบ PromptPay/CRC", "src/PromptPayQR.cpp", "tests/test_promptpay.cpp และสแกนด้วยแอปธนาคาร"),
        ("หน้าเว็บ/API บน ESP32", "src/main.cpp: handleRoot/handlers", "curl ทุก endpoint และตรวจ OLED/Serial"),
        ("GPIO/OLED", "src/main.cpp: constants/render functions", "บอร์ดจริงและสายจริง"),
        ("WiFi/NVS", "connectWifi(), handleSetup(), handleConfig()", "ปิดเปิดไฟและเปลี่ยน Router"),
        ("Dashboard", "local-dashboard/server.js", "npm run test:dashboard"),
        ("Payment Gateway", "local-dashboard/omise.js หรือ adapter ใหม่", "test adapter, webhook signature และ callback จริง"),
    ], [2500, 3900, 2960], font_size=8.5)
    add_code_block(doc, "g++ -std=c++17 -Iinclude tests/test_promptpay.cpp src/PromptPayQR.cpp -o /tmp/paymentesp-test\n/tmp/paymentesp-test\nnpm run test:dashboard\nnpm run test:omise\npio run -e esp32dev")
    add_callout(doc, "Definition of Done", "แก้โค้ดแล้วต้อง Build ผ่าน รัน automated tests และทดสอบ Hardware Acceptance บนบอร์ด/OLED ตัวจริงอีกครั้ง", "success")

    add_heading(doc, "4. Build และ Upload Firmware", 1)
    add_heading(doc, "4.1 Checklist ก่อน Build", 2)
    add_list(doc, [
        "VS Code เปิดที่โฟลเดอร์ source/ และเห็น platformio.ini",
        "เลือก Environment esp32dev",
        "ต่อ OLED ตาม GPIO 21/22 และตรวจ VCC/GND",
        "ต่อ ESP32 ด้วยสาย USB Data",
        "ปิด Serial Monitor หรือโปรแกรมอื่นที่กำลังจับ Port",
    ])
    add_heading(doc, "4.2 Build Hardware Environment", 2)
    add_code_block(doc, "pio run -e esp32dev")
    add_table(doc, ["ผล Build รอบส่งมอบ", "ค่า"], [
        ("PlatformIO Core", "6.1.19"),
        ("RAM", "50,408 / 327,680 bytes (15.4%)"),
        ("Flash", "1,052,917 / 1,310,720 bytes (80.3%)"),
        ("Result", "SUCCESS"),
    ], [3600, 5760], font_size=10)
    add_body(doc, "ใน VS Code สามารถใช้ PlatformIO sidebar > Project Tasks > esp32dev > General > Build ได้ผลเท่ากับคำสั่งด้านบน ผู้ที่ไม่ใช้ IDE สามารถรันคำสั่ง pio จาก PlatformIO Core CLI ได้")
    add_heading(doc, "4.3 Upload", 2)
    add_list(doc, [
        "ต่อ ESP32 ด้วยสาย USB Data",
        "ปิด Serial Monitor หรือโปรแกรมอื่นที่จับ Port",
        "รันคำสั่ง Upload",
        "ถ้าค้างที่ Connecting ให้กด BOOT ค้างและปล่อยเมื่อเริ่มเขียน Flash",
        "รอข้อความ SUCCESS และห้ามถอดสายระหว่างเขียน Flash",
    ], ordered=True)
    add_code_block(doc, "pio run -e esp32dev -t upload")
    add_body(doc, "ถ้ามีหลายบอร์ด ให้ระบุ Port ชัดเจนเพื่อป้องกัน Upload ผิดเครื่อง")
    add_code_block(doc, "pio run -e esp32dev -t upload --upload-port <PORT>")
    add_heading(doc, "4.4 เปิด Serial Monitor", 2)
    add_code_block(doc, "pio device monitor -b 115200")
    add_callout(doc, "ผลที่คาดหวัง", "Serial แสดง PaymentESP, Setup AP หรือ WiFi connected, IP address และ HTTP server started on port 80", "success")
    add_heading(doc, "4.5 ล้างค่าเดิมและ Factory Reset", 2)
    add_body(doc, "ใช้เมื่อบอร์ดมี WiFi/PromptPay เดิมจากการทดสอบ หรือต้องส่งเครื่องใหม่ให้ลูกค้า คำสั่ง erase จะล้างทั้ง Firmware และ NVS จึงต้อง Upload ใหม่ทันที")
    add_code_block(doc, "pio run -e esp32dev -t erase\npio run -e esp32dev -t upload")
    add_callout(doc, "คำเตือน", "ห้าม Factory Reset เครื่องลูกค้าที่ใช้งานอยู่โดยไม่สำรองค่าหน้างาน เพราะ WiFi, PromptPay ID และ Webhook จะหาย", "danger")
    add_heading(doc, "4.6 ทางเลือก: Flash ไฟล์รวม", 2)
    add_body(doc, "โฟลเดอร์ firmware มี firmware-merged.bin สำหรับแฟลชที่ offset 0x0 เหมาะกับเครื่องที่ติดตั้ง esptool แล้ว โดยวิธี PlatformIO ด้านบนยังเป็นวิธีแนะนำ")
    add_code_block(doc, "python -m esptool --chip esp32 --port <PORT> --baud 460800 write_flash 0x0 firmware/firmware-merged.bin")
    add_heading(doc, "4.7 ตรวจว่า Upload สำเร็จ", 2)
    add_table(doc, ["จุดตรวจ", "ผลที่ต้องเห็น"], [
        ("PlatformIO", "SUCCESS และไม่มี upload error"),
        ("Serial 115200", "ESP32 PromptPay QR Payment และ HTTP server started on port 80"),
        ("OLED", "PaymentESP SETUP หรือ PaymentESP READY"),
        ("Setup AP", "พบ PaymentESP-Setup หากยังไม่มี WiFi ที่ใช้ได้"),
    ], [3100, 6260], font_size=9.5)

    doc.add_page_break()
    add_heading(doc, "5. ตั้งค่าเครื่องครั้งแรก", 1)
    add_heading(doc, "5.1 สถานะที่ OLED หลัง Boot", 2)
    add_table(doc, ["สถานะ", "ข้อความที่แสดง", "การทำต่อ"], [
        ("SETUP", "PaymentESP SETUP / PaymentESP-Setup / 192.168.4.1/setup", "เชื่อม Setup AP แล้วกรอกค่าหน้างาน"),
        ("READY", "PaymentESP READY / IP / paymentesp.local", "เปิดหน้าเว็บด้วย .local หรือ IP"),
        ("OLED ว่าง", "ไม่มีข้อความ", "ตรวจสาย, 3V3, GPIO 21/22 และ address 0x3C"),
    ], [1800, 4700, 2860], font_size=8.9)
    add_heading(doc, "5.2 Setup ผ่าน Access Point", 2)
    add_list(doc, [
        "จ่ายไฟให้ ESP32 และรอประมาณ 20 วินาที",
        "เชื่อม WiFi ชื่อ PaymentESP-Setup ด้วยรหัส paymentesp",
        "หากโทรศัพท์แจ้งว่าไม่มี Internet ให้เลือกเชื่อมต่อต่อไป",
        "เปิด Browser ที่ http://192.168.4.1/setup",
        "กรอก WiFi SSID/Password ของสถานที่ใช้งาน",
        "กรอก PromptPay ID และ Webhook URL ถ้ามี",
        "กด Save and restart ESP32",
    ], ordered=True)
    add_callout(doc, "โทรศัพท์ไม่มี Internet", "ขณะต่อ PaymentESP-Setup โทรศัพท์อาจแจ้งว่าเครือข่ายไม่มี Internet ให้เลือก Keep connection/เชื่อมต่อต่อไป แล้วเปิด URL ด้วยตัวเอง", "note")
    add_heading(doc, "5.3 รายละเอียดช่องตั้งค่า", 2)
    add_table(doc, ["ช่อง", "ตัวอย่าง", "ข้อกำหนด"], [
        ("WiFi SSID", "SHOP-WIFI", "2.4 GHz, ไม่เกิน 32 ตัวอักษร"),
        ("WiFi Password", "********", "ไม่เกิน 63 ตัวอักษร"),
        ("PromptPay ID", "0812345678", "มือถือ 10 หลัก, ID 13 หลัก หรือ e-wallet 15 หลัก"),
        ("Webhook URL", "http://192.168.1.10:3001/api/webhook", "เว้นว่างได้; ห้ามใช้ localhost บน ESP32"),
    ], [2200, 3100, 4060], font_size=9)
    add_callout(doc, "การเก็บค่า", "WiFi, PromptPay ID และ Webhook ถูกเก็บใน NVS ของ ESP32 และยังอยู่หลังปิดเปิดเครื่อง", "note")
    add_heading(doc, "5.4 ทดสอบ NVS หลังตั้งค่า", 2)
    add_list(doc, [
        "รอให้บอร์ดเชื่อม WiFi และ OLED แสดง READY/IP",
        "ถอดสาย USB หรือแหล่งจ่ายไฟอย่างน้อย 5 วินาที",
        "จ่ายไฟใหม่และตรวจว่าบอร์ดเชื่อม WiFi เดิมโดยไม่เปิด Setup Portal",
        "เปิด /api/config และตรวจ PromptPay ID/Webhook โดยไม่แสดง WiFi Password",
    ], ordered=True)
    add_heading(doc, "5.5 เมื่อย้ายสถานที่หรือ WiFi เปลี่ยน", 2)
    add_body(doc, "หาก ESP32 เชื่อม WiFi เดิมไม่ได้ภายในประมาณ 20 วินาที ระบบจะเปิด PaymentESP-Setup อัตโนมัติ ให้เชื่อม AP แล้วบันทึก WiFi ใหม่ ค่า PromptPay/Webhook เดิมสามารถแก้พร้อมกันได้")

    add_heading(doc, "6. เข้าใช้งานบน WiFi จริง", 1)
    add_list(doc, [
        "เชื่อมโทรศัพท์หรือคอมพิวเตอร์กลับเข้า WiFi เดียวกับ ESP32",
        "ดู OLED หรือ Serial เพื่ออ่าน IP ที่ได้รับ",
        "เปิด http://paymentesp.local",
        "ถ้า .local ไม่ทำงาน ให้เปิดด้วย IP เช่น http://192.168.1.50",
        "ตรวจ GET /api/config ว่า wifiConnected เป็น true และ setupMode เป็น false",
    ], ordered=True)
    add_code_block(doc, "curl http://paymentesp.local/api/config")
    add_code_block(doc, "{\n  \"wifiConnected\": true,\n  \"setupMode\": false,\n  \"ip\": \"192.168.1.50\",\n  \"hostname\": \"http://paymentesp.local\"\n}")
    add_callout(doc, "Network", "อุปกรณ์ที่เปิดหน้าเว็บต้องอยู่ในวง LAN เดียวกับ ESP32 และ Router ต้องไม่เปิด Client Isolation", "warning")
    add_heading(doc, "6.1 ถ้า paymentesp.local ใช้ไม่ได้", 2)
    add_list(doc, [
        "อ่าน IP จาก OLED หรือ Serial Monitor",
        "เปิด http://<ESP32-IP> เช่น http://192.168.1.50",
        "ตรวจว่าโทรศัพท์/คอมพิวเตอร์อยู่ WiFi/VLAN เดียวกับ ESP32",
        "ปิด VPN ชั่วคราวและตรวจ Client/AP Isolation ของ Router",
    ], ordered=True)

    add_heading(doc, "7. Local Dashboard สำหรับเครื่องจริง", 1)
    add_heading(doc, "7.1 ติดตั้งและรัน Dashboard จาก Source", 2)
    add_body(doc, "Dashboard รันบนคอมพิวเตอร์ ไม่ได้รันบน ESP32 ให้เปิด Terminal ที่โฟลเดอร์ source/ แล้วตรวจ Node.js ก่อนติดตั้ง package")
    add_code_block(doc, "node --version\nnpm --version\nnpm ci\nnpm run dashboard")
    add_body(doc, "เปิด http://localhost:3001 แล้วตั้ง ESP32 Hardware URL เป็น http://paymentesp.local หรือ IP ของบอร์ด จากนั้นเปิด Sync to ESP32 hardware")
    add_heading(doc, "7.2 ตั้งค่า Dashboard ให้คุยกับ Hardware", 2)
    add_table(doc, ["ช่อง", "ค่าที่แนะนำ", "หมายเหตุ"], [
        ("ESP32 Hardware URL", "http://paymentesp.local", "ใช้ IP หาก mDNS ไม่ทำงาน"),
        ("PromptPay ID", "หมายเลขผู้รับจริง", "กด Save Settings เพื่อบันทึกลง NVS"),
        ("Webhook URL", "http://<PC-IP>:3001/api/webhook", "PC ต้องเปิด Firewall port 3001"),
        ("Amount", "15.00", "ใช้เฉพาะ Dynamic QR"),
        ("Reference", "ORDER-0001", "ใช้ติดตาม Logs"),
    ], [2350, 2800, 4210], font_size=9.2)
    add_callout(doc, "Webhook", "ESP32 มอง localhost เป็นตัว ESP32 เอง จึงต้องใช้ LAN IP ของคอมพิวเตอร์หรือ Server ที่รับ POST", "note")
    add_heading(doc, "7.3 หา LAN IP ของคอมพิวเตอร์", 2)
    add_table(doc, ["ระบบ", "คำสั่ง/ตำแหน่ง"], [
        ("Windows", "ipconfig แล้วดู IPv4 Address ของ WiFi"),
        ("macOS", "System Settings > Wi-Fi > Details หรือ ipconfig getifaddr en0"),
        ("Linux", "hostname -I หรือ ip addr"),
    ], [2500, 6860], font_size=9.5)
    add_heading(doc, "7.4 ใช้งานแบบไม่เปิด Dashboard", 2)
    add_body(doc, "ESP32 มีหน้าเว็บและ API ในตัว ลูกค้าสามารถเปิด paymentesp.local หรือ IP เพื่อสร้าง QR ได้โดยไม่ต้องเปิดคอมพิวเตอร์ Dashboard แต่ REST receiver บน Dashboard จะไม่ทำงานเมื่อคอมพิวเตอร์ปิด")

    add_heading(doc, "8. สร้างและทดสอบ QR", 1)
    add_heading(doc, "8.1 Static QR", 2)
    add_body(doc, "Static QR ไม่มี Tag ยอดเงิน ผู้สแกนสามารถกรอกยอดในแอปธนาคารได้เอง")
    add_code_block(doc, "curl \"http://paymentesp.local/api/static?ref=STATIC-TEST\"")
    add_heading(doc, "8.2 Dynamic QR", 2)
    add_body(doc, "Dynamic QR ฝังยอดเงินไว้ใน QR แอปธนาคารจะแสดงยอดที่ระบบกำหนดและโดยทั่วไปไม่ให้ผู้จ่ายแก้ไข")
    add_code_block(doc, "curl \"http://paymentesp.local/api/dynamic?amount=15.00&ref=DYNAMIC-TEST\"")
    add_heading(doc, "8.3 ตรวจด้วยแอปธนาคาร", 2)
    add_list(doc, [
        "สแกน QR จาก OLED หรือหน้าเว็บ",
        "ตรวจชื่อผู้รับให้ตรงกับเจ้าของ PromptPay",
        "ตรวจยอดสำหรับ Dynamic QR",
        "ยกเลิกก่อนยืนยันหากเป็นการทดสอบที่ไม่ต้องการโอนเงินจริง",
    ], ordered=True)
    add_callout(doc, "ข้อจำกัด", "การสแกนหรือโอนเงินจริงไม่ทำให้ ESP32 รู้ผลโดยอัตโนมัติ หากไม่มี Bank API/Payment Gateway ส่ง callback", "danger")

    add_heading(doc, "9. Serial, REST POST และ Logs", 1)
    add_code_block(doc, "=== QR Payment Event ===\nMode: dynamic\nAmount: 15.00\nReference: ORDER-0001\nWebhook status: 202")
    add_table(doc, ["Output", "ตรวจที่", "ความหมาย"], [
        ("OLED", "หน้าจอเครื่อง", "QR, mode, amount, reference"),
        ("Serial", "Baud 115200", "รายละเอียด Event และสถานะ Network"),
        ("Device Logs", "GET /api/logs", "20 รายการล่าสุดใน RAM"),
        ("REST POST", "Webhook receiver", "Event JSON หลังสร้าง QR"),
    ], [2100, 2600, 4660], font_size=9.5)
    add_callout(doc, "อย่าสับสน", "HTTP 202 และ Device Logs ยืนยันว่า Event การสร้าง QR ถูกส่ง/รับแล้ว ไม่ใช่หลักฐานว่าเงินเข้าบัญชี", "warning")

    add_heading(doc, "10. API Reference", 1)
    add_table(doc, ["Method", "Endpoint", "หน้าที่", "Parameter"], [
        ("GET", "/", "หน้า QR ของ ESP32", "ไม่มี"),
        ("GET/POST", "/setup", "ตั้งค่า WiFi/PromptPay", "ssid, password, promptpay, webhook"),
        ("GET", "/api/static", "สร้าง Static QR", "ref"),
        ("GET", "/api/dynamic", "สร้าง Dynamic QR", "amount, ref"),
        ("GET", "/api/logs", "ดู Device Logs", "ไม่มี"),
        ("GET/POST", "/api/config", "ดู/บันทึก PromptPay และ Webhook", "promptpay, webhook"),
        ("POST", "/api/payment", "รับสถานะ paid จากระบบที่เชื่อถือได้", "paymentId, status, amount, ref"),
    ], [1200, 2200, 3100, 2860], font_size=8.6)

    add_heading(doc, "11. Troubleshooting", 1)
    add_table(doc, ["อาการ", "สาเหตุ", "วิธีแก้"], [
        ("PlatformIO ไม่พบ Project", "เปิดโฟลเดอร์ผิดระดับ", "เปิด source/ ที่มี platformio.ini"),
        ("pio: command not found", "CLI ไม่อยู่ใน PATH", "ใช้ PlatformIO Project Tasks หรือ Terminal ของ extension"),
        ("Build ดาวน์โหลด library ไม่ได้", "Internet/Proxy/Firewall", "ตรวจ Network แล้วรัน pio pkg install -e esp32dev"),
        ("ไม่พบ Serial Port", "สายชาร์จอย่างเดียว/ไม่มี Driver", "เปลี่ยนสาย Data และติดตั้ง CP210x/CH340"),
        ("Permission denied ที่ Serial", "สิทธิ์ Port ไม่พอ", "ปิดโปรแกรมที่จับ Port; Linux เพิ่มกลุ่ม dialout"),
        ("Upload ค้าง Connecting", "บอร์ดไม่เข้า Bootloader", "กด BOOT ค้างระหว่าง Connecting"),
        ("OLED ไม่แสดง", "สาย/address ผิด", "ตรวจ GPIO 21/22, 3V3 และ 0x3C"),
        ("ไม่พบ PaymentESP-Setup", "ESP32 เชื่อม WiFi สำเร็จแล้ว", "ดู IP บน OLED หรือปิด Routerชั่วคราวแล้วรีสตาร์ต"),
        ("paymentesp.local เปิดไม่ได้", "mDNS ไม่รองรับ", "ใช้ IP จาก OLED/Serial"),
        ("Dashboard timeout", "Device URL/IP ผิดหรือคนละ LAN", "แก้ URL และปิด Client Isolation"),
        ("REST POST ว่าง", "ใช้ localhost หรือ Firewall บล็อก", "ใช้ PC LAN IP และเปิด port 3001"),
        ("WiFi เปลี่ยน", "ค่าที่บันทึกเชื่อมไม่ได้", "รอ AP fallback แล้วตั้งที่ 192.168.4.1/setup"),
    ], [2500, 3300, 3560], font_size=8.6)

    add_heading(doc, "12. Hardware Acceptance ก่อนส่งลูกค้า", 1)
    add_list(doc, [
        "Upload environment esp32dev ลงบอร์ดตัวที่จะส่งสำเร็จ",
        "ทดสอบ Setup AP และหน้า 192.168.4.1/setup",
        "เชื่อม WiFi จริงและ OLED แสดง READY/IP",
        "เปิด paymentesp.local หรือ IP ได้จากโทรศัพท์ในวงเดียวกัน",
        "บันทึก PromptPay ID แล้วปิดเปิดเครื่อง ค่ายังคงเดิม",
        "Static QR สแกนได้และไม่มียอดฝัง",
        "Dynamic QR สแกนได้และยอดตรงตามที่กำหนด",
        "ตรวจชื่อผู้รับในแอปธนาคารถูกต้อง",
        "Serial, Device Logs และ REST POST มี reference/amount ตรงกัน",
        "ติดป้าย SSID Setup, Password และวิธีกู้คืนไว้กับเครื่อง",
    ], ordered=True)
    add_callout(doc, "เกณฑ์ส่งมอบ", "ต้องผ่านครบทุกข้อบนบอร์ดและ OLED ตัวจริงก่อนส่งให้ลูกค้า ภาพ Wokwi ใช้แทน Hardware Acceptance ไม่ได้", "success")

    add_heading(doc, "13. ขั้นตอนส่งมอบเครื่องให้ลูกค้า", 1)
    add_heading(doc, "13.1 เตรียมบอร์ดก่อนปิดกล่อง", 2)
    add_list(doc, [
        "ใช้บอร์ด ESP32 และ OLED ตัวเดียวกับที่ผ่าน Hardware Acceptance",
        "ตรวจสายทุกจุดและยึดสาย/ขั้วต่อไม่ให้หลุดระหว่างขนส่ง",
        "ปิดข้อมูลทดสอบเดิม แล้วตั้ง PromptPay ID ของลูกค้าผ่าน Setup Portal",
        "ทดสอบปิดเปิดไฟโดยไม่ต่อคอมพิวเตอร์อย่างน้อย 2 รอบ",
        "บันทึก Asset ID, MAC address, Firmware version และวันที่ทดสอบ",
        "ถ่ายหลักฐาน OLED READY, Dynamic QR, Serial, Device Logs และ REST POST",
    ], ordered=True)
    add_heading(doc, "13.2 ข้อมูลที่ให้ลูกค้า", 2)
    add_table(doc, ["รายการ", "ค่าที่ส่งมอบ/กรอกหน้างาน"], [
        ("Setup AP SSID", "PaymentESP-Setup"),
        ("Setup AP Password", "paymentesp"),
        ("Setup URL", "http://192.168.4.1/setup"),
        ("ใช้งานปกติ", "http://paymentesp.local หรือ IP จาก OLED"),
        ("ไฟเลี้ยง", "USB 5V คุณภาพดี; บอร์ดควบคุม OLED ที่ 3.3V"),
        ("Serial", "115200 baud สำหรับช่าง"),
    ], [3200, 6160], font_size=9.5)
    add_heading(doc, "13.3 ไฟล์ที่ต้องส่ง", 2)
    add_list(doc, [
        "ZIP ชุดส่งมอบที่มี source/, firmware/, documents/ และ SHA256SUMS.txt",
        "คู่มือติดตั้ง Hardware ฉบับนี้",
        "Wokwi Simulator Test Report และ Hardware Test Report",
        "ภาพหลักฐานจากเครื่องจริงหลังปิดบังข้อมูลสำคัญ",
        "ข้อมูล WiFi Password/Secret สำหรับหน้างานให้ส่งแยกช่องทาง ไม่ใส่ใน ZIP",
    ])
    add_heading(doc, "13.4 Quick Start สำหรับลูกค้า", 2)
    add_list(doc, [
        "จ่ายไฟให้เครื่องและรอ OLED แสดง READY",
        "เปิด paymentesp.local หรือ IP จาก OLED",
        "เลือก Static QR หรือ Dynamic QR แล้วกรอกยอดเงิน",
        "ตรวจชื่อผู้รับและยอดเงินในแอปธนาคารก่อนยืนยัน",
        "หาก WiFi เปลี่ยน รอ PaymentESP-Setup แล้วตั้งค่าใหม่ที่ 192.168.4.1/setup",
    ], ordered=True)
    add_callout(doc, "บริการหลังส่งมอบ", "ก่อนแก้ Source หรือ Factory Reset ให้สำรองค่าหน้างานและบันทึก Firmware version ทุกครั้ง", "note")

    add_heading(doc, "14. Wokwi สำหรับทดสอบเสริม", 1)
    add_body(doc, "Wokwi แยก environment ออกจาก Hardware เพื่อไม่ให้ค่า Wokwi-GUEST และ host.wokwi.internal ติดไปใน firmware ที่ส่งลูกค้า")
    add_code_block(doc, "pio run -e wokwi")
    add_list(doc, [
        "เปิด diagram.json",
        "สั่ง Wokwi: Start Simulator",
        "เปิด localhost:8180 หลัง Simulator boot",
        "ใช้เฉพาะตรวจ logic หรือสาธิตก่อนมี Hardware",
    ], ordered=True)

    doc.add_page_break()
    add_heading(doc, "ภาคผนวก A: หลักฐานจากเครื่องจริง", 1)
    add_callout(doc, "ต้องจัดทำก่อนส่ง", "หลักฐานในตารางนี้ต้องถ่ายจาก ESP32/OLED ตัวที่จะส่งลูกค้า หลัง Upload และตั้งค่าครบแล้ว ไม่ใช้ภาพ Wokwi แทน", "warning")
    add_table(doc, ["ไฟล์หลักฐาน", "สิ่งที่ต้องเห็น", "สถานะรอบจัดทำ"], [
        ("01-hardware-wiring.jpg", "ESP32, OLED และสาย GPIO 21/22/3V3/GND", "รอถ่ายจากเครื่องจริง"),
        ("02-oled-ready.jpg", "OLED แสดง PaymentESP READY และ IP", "รอถ่ายจากเครื่องจริง"),
        ("03-oled-dynamic-qr.jpg", "Dynamic QR, amount และ reference", "รอถ่ายจากเครื่องจริง"),
        ("04-serial-output.png", "WiFi connected, IP และ QR Event", "รอจับภาพจากเครื่องจริง"),
        ("05-device-logs.png", "Device Logs มี amount/reference", "รอจับภาพจากเครื่องจริง"),
        ("06-rest-post.png", "Webhook รับ Event และ HTTP status", "รอจับภาพจากเครื่องจริง"),
    ], [3000, 4000, 2360], font_size=8.8)
    add_body(doc, "หลังจัดทำหลักฐาน ให้ตรวจว่า PromptPay ID, WiFi Password และข้อมูลส่วนตัวที่ไม่จำเป็นถูกปิดบังก่อนส่งไฟล์ให้ลูกค้า")

    add_heading(doc, "ภาคผนวก B: ชุดไฟล์ส่งมอบ", 1)
    add_table(doc, ["รายการ", "ตำแหน่ง"], [
        ("คู่มือ Hardware", "documents/PaymentESP-Hardware-Installation-Manual-TH.docx"),
        ("Hardware Test Report", "documents/PaymentESP-Hardware-Test-Report-TH.docx"),
        ("Simulator Test Report", "documents/PaymentESP-Wokwi-Simulator-Test-Report-TH.docx"),
        ("Requirement Checklist", "documents/requirements-checklist.md"),
        ("Source Code", "source/"),
        ("Firmware Hardware", "firmware/"),
        ("หลักฐาน", "evidence/"),
        ("ตรวจความถูกต้อง", "SHA256SUMS.txt"),
    ], [3000, 6360], font_size=10)
    add_body(doc, "หมายเหตุ: อย่าใส่ WiFi Password, ไฟล์ .env หรือ Secret Key จริงลงใน ZIP ชุดส่งมอบ", bold_prefix="หมายเหตุ:")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
