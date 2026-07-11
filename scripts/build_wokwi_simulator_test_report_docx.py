#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_installation_docx import (
    BLACK,
    GREEN,
    MUTED,
    NAVY,
    add_body,
    add_callout,
    add_code_block,
    add_figure,
    add_heading,
    set_run_font,
)
from build_hardware_test_report_docx import (
    add_report_list,
    add_report_table,
    configure_report,
)


ROOT = Path(__file__).resolve().parents[1]
PACKAGE_NAME = "PaymentESP-v1.1.0-Hardware-20260628"
OUTPUT = ROOT / "deliverables" / PACKAGE_NAME / "documents" / "PaymentESP-Wokwi-Simulator-Test-Report-TH.docx"
EVIDENCE_DIR = ROOT / "docs" / "evidence"
LOGS_IMAGE = EVIDENCE_DIR / "03-device-logs.png"
REST_IMAGE = EVIDENCE_DIR / "04-rest-post-events.png"
TEST_TIME = "27 มิถุนายน 2026 เวลา 20:46 น. (Asia/Bangkok)"


def add_masthead(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(20)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    set_run_font(kicker.add_run("WOKWI SIMULATOR TEST EVIDENCE"), size=10.5, color=GREEN, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    set_run_font(title.add_run("รายงานผลการทดสอบบน Wokwi Simulator"), size=24, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("PaymentESP | ESP32 PromptPay QR + OLED + REST POST"), size=13, color=MUTED)

    metadata = [
        ("รอบทดสอบ", "EVIDENCE-001"),
        ("เวลาทดสอบ", TEST_TIME),
        ("อุปกรณ์จำลอง", "ESP32 DevKit V1 และ OLED SSD1306 128x64"),
        ("รูปแบบ QR", "Dynamic QR"),
        ("ยอดทดสอบ", "15.00 บาท"),
        ("สถานะ", "SIMULATOR TEST PASS"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        set_run_font(paragraph.add_run(f"{label}: "), size=10.5, color=BLACK, bold=True)
        set_run_font(paragraph.add_run(value), size=10.5, color=BLACK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(14)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GREEN)
    p_bdr.append(bottom)
    rule._p.get_or_add_pPr().append(p_bdr)

    add_callout(
        doc,
        "สรุปผล",
        "Wokwi สร้าง Dynamic QR ยอด 15.00 บาท แสดงสถานะบน OLED จำลอง ส่ง Serial event และ REST POST ได้ HTTP 202 โดย reference ตรงกัน",
        "success",
    )


def build_document():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_report(doc)
    header = doc.sections[0].header.paragraphs[0]
    header.clear()
    set_run_font(header.add_run("PaymentESP | Wokwi Simulator Test Report"), size=8.5, color=MUTED, bold=True)

    props = doc.core_properties
    props.title = "รายงานผลการทดสอบ PaymentESP บน Wokwi Simulator"
    props.subject = "ESP32 PromptPay QR simulator test evidence"
    props.author = "PaymentESP Project Team"
    props.keywords = "ESP32, Wokwi, PromptPay, OLED, Serial, REST POST, Logs"
    props.comments = "Simulator evidence report; sensitive PromptPay data redacted"

    add_masthead(doc)

    add_heading(doc, "1. วัตถุประสงค์", 1)
    add_body(doc, "เอกสารฉบับนี้บันทึกสิ่งที่ทดสอบบน Wokwi Simulator และ Local Dashboard ในรอบ EVIDENCE-001 โดยเน้นการสร้าง QR, การแสดงผล OLED, Serial Output, REST POST และ Logging ตาม Requirement ของโครงการ")
    add_callout(doc, "ขอบเขต", "รายงานนี้เป็นผลจาก Simulator ไม่ใช่ผลตรวจ ESP32/OLED เครื่องจริง และ HTTP 202 ไม่ใช่การยืนยันว่าเงินเข้าบัญชี", "warning")

    doc.add_page_break()
    add_heading(doc, "2. Test Setup", 1)
    add_report_table(doc, ["รายการ", "ค่าที่ใช้ทดสอบ"], [
        ("Firmware environment", "wokwi"),
        ("Board", "ESP32 DevKit V1 จำลอง"),
        ("Display", "OLED SSD1306 128x64 I2C"),
        ("Dashboard", "Local Dashboard ที่ localhost:3001"),
        ("QR mode", "Dynamic"),
        ("Amount", "15.00 บาท"),
        ("Reference", "EVIDENCE-001"),
        ("Webhook", "Local REST receiver /api/webhook"),
    ], [3000, 6360], font_size=9.5)

    add_heading(doc, "3. ขั้นตอนที่ทดสอบ", 1)
    add_report_list(doc, [
        "Build Firmware ด้วย environment wokwi",
        "เริ่ม Wokwi Simulator และตรวจว่า HTTP server เปิดที่ port 80",
        "สร้าง Dynamic PromptPay QR ด้วยยอด 15.00 บาทและ reference EVIDENCE-001",
        "ตรวจข้อความและ QR state บน OLED จำลอง",
        "ตรวจ Serial event ที่ baud rate 115200",
        "ตรวจ GET /api/logs ว่าบันทึกเวลา ยอดเงิน reference และสถานะ POST",
        "ตรวจ REST receiver ว่าได้รับ event เดียวกันและตอบ HTTP 202",
    ], ordered=True)

    doc.add_page_break()
    add_heading(doc, "4. ผลการทดสอบ Simulator", 1)
    add_report_table(doc, ["ID", "รายการทดสอบ", "ผลที่พบ", "สถานะ"], [
        ("SIM-01", "Wokwi firmware build", "Compile environment wokwi สำเร็จ", "PASS"),
        ("SIM-02", "HTTP server startup", "Serial แสดง HTTP server started on port 80", "PASS"),
        ("SIM-03", "Dynamic QR generation", "Mode dynamic, amount 15.00, reference EVIDENCE-001", "PASS"),
        ("SIM-04", "OLED display", "แสดง PromptPay, Dynamic, 15.00 และ reference", "PASS"),
        ("SIM-05", "Serial Output", "Event มี mode, amount, reference, time และ webhook status", "PASS"),
        ("SIM-06", "Device Logs", "GET /api/logs พบรายการยอด 15.00 และ POST 202", "PASS"),
        ("SIM-07", "REST POST", "POST /api/webhook ได้ HTTP 202 Accepted", "PASS"),
        ("SIM-08", "Reference matching", "Device log และ REST receiver ตรงกันที่ EVIDENCE-001", "PASS"),
        ("SIM-09", "Timestamp logging", "บันทึกเวลา createdAt และ receivedAt ได้", "PASS"),
    ], [1100, 2700, 4010, 1550], font_size=8.5)
    add_callout(doc, "ผลรวม", "Simulator Test ผ่าน 9 รายการ, Fail 0 รายการ", "note")

    add_heading(doc, "5. OLED Simulator Result", 1)
    add_report_table(doc, ["Field", "ค่าที่แสดงบน OLED จำลอง"], [
        ("Title", "PromptPay"),
        ("Mode", "Dynamic"),
        ("Amount", "15.00"),
        ("Reference", "EVIDENCE-001"),
        ("QR", "สร้างและแสดงบน OLED canvas สำเร็จ"),
    ], [2700, 6660], font_size=9.5)
    add_callout(doc, "การปกป้องข้อมูล", "ไม่แนบภาพ QR ต้นฉบับในเอกสารลูกค้า เพราะ QR มี PromptPay target จริงและสามารถนำไปสแกนได้", "warning")

    add_heading(doc, "6. Serial Output (Sanitized)", 1)
    add_code_block(doc, "ESP32 PromptPay QR Payment\nHTTP server started on port 80\nREST POST status: 202\n\n=== QR Payment Event ===\nMode: dynamic\nPromptPay ID: [REDACTED]\nAmount: 15.00\nReference: EVIDENCE-001\nCreated at: 2026-06-27T20:46:39+07:00\nWebhook status: 202\nPayload: [REDACTED]")
    add_body(doc, "Serial หลักฐานต้นฉบับถูกเก็บไว้ภายในโครงการ แต่รายงานฉบับส่งมอบปิด PromptPay ID และ payload เพื่อป้องกันข้อมูลการชำระเงิน")

    add_heading(doc, "7. Device Logs Evidence", 1)
    add_figure(doc, LOGS_IMAGE, "ภาพที่ 1 Device Logs จาก Wokwi/ESP32 API: Dynamic 15.00 และ POST 202", width=5.1)
    add_body(doc, "รายการ EVIDENCE-001 แสดง mode dynamic, amount 15.00, เวลา 2026-06-27T20:46:39+07:00 และสถานะ POST 202 สอดคล้องกับ Serial event")

    add_heading(doc, "8. REST POST Evidence", 1)
    add_figure(doc, REST_IMAGE, "ภาพที่ 2 REST receiver: HTTP 202, matched reference EVIDENCE-001 และ amount 15.00", width=5.1)
    add_body(doc, "REST receiver บันทึก createdAt จาก ESP32 และ receivedAt ฝั่ง Dashboard พร้อมจับคู่ reference EVIDENCE-001 ได้สำเร็จ")

    add_heading(doc, "9. ข้อสรุปและข้อจำกัด", 1)
    add_report_list(doc, [
        "Wokwi ยืนยัน flow การสร้าง Dynamic QR, OLED state, Serial, REST POST และ Logs ได้",
        "ยอดเงิน 15.00 และ reference EVIDENCE-001 ตรงกันตลอด flow",
        "HTTP 202 หมายถึง receiver รับ event แล้ว ไม่ได้ตรวจสอบเงินเข้าบัญชี",
        "Simulator ไม่ยืนยัน USB upload, สาย OLED, WiFi/NVS หรือไฟเลี้ยงของบอร์ดจริง จึงต้องทำ Hardware Acceptance แยกก่อนส่งลูกค้า",
    ])

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
